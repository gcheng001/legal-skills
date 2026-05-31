#!/usr/bin/env python3
"""
材料分析脚本 - v4.0 Markdown优先版
先生成Markdown报告（可校验），再转换为Word文档
"""

import os
import re
import json
from pathlib import Path
from datetime import datetime
from PIL import Image
import subprocess

# 导入PDF截图模块
import sys
sys.path.insert(0, str(Path(__file__).parent))
from extract_pdf_screenshot import extract_screenshots_for_evidence


def scan_evidence_files(project_path):
    """
    智能扫描所有证据文件
    返回按证据类型分类的文件列表
    """
    project_path = Path(project_path)
    evidence_files = {
        '中标公告': [],
        '支付协议': [],
        '银行回单': [],
        '结算单': [],
        '其他': []
    }

    # 扫描所有文件
    all_files = []
    for file_path in project_path.rglob('*'):
        if file_path.is_file() and not any(x in str(file_path) for x in ['立案材料', '截图证据', '.temp_', '材料分析报告', 'markdown']):
            if file_path.suffix.lower() in ['.pdf', '.png', '.jpg', '.jpeg', '.doc', '.docx']:
                all_files.append(file_path)

    # 先按文件名分类
    unclassified = []
    for file_path in all_files:
        name_lower = file_path.name.lower()

        if '中标' in name_lower or '公告' in name_lower:
            evidence_files['中标公告'].append(file_path)
        elif '支付协议' in name_lower or '付款协议' in name_lower:
            evidence_files['支付协议'].append(file_path)
        elif '银行' in name_lower or '回单' in name_lower:
            evidence_files['银行回单'].append(file_path)
        elif '结算' in name_lower:
            evidence_files['结算单'].append(file_path)
        else:
            unclassified.append(file_path)

    # 智能分类未识别的图片（根据图片特征）
    bank_images = []
    doc_images = []

    for file_path in unclassified:
        if file_path.suffix.lower() in ['.png', '.jpg', '.jpeg']:
            try:
                img = Image.open(file_path)
                width, height = img.size
                ratio = width / height

                # 银行回单特征：横版，宽度>高度，比例>1.5
                if ratio > 1.5:
                    bank_images.append(file_path)
                # 文档照片特征：竖版，高度>宽度，比例<1.0
                elif ratio < 1.0:
                    doc_images.append(file_path)
                else:
                    evidence_files['其他'].append(file_path)
            except:
                evidence_files['其他'].append(file_path)
        else:
            evidence_files['其他'].append(file_path)

    # 分配银行回单（2个）
    if len(bank_images) >= 2:
        evidence_files['银行回单'] = bank_images[:2]

    # 分配文档（支付协议和结算单）
    if len(doc_images) >= 2:
        # 按文件大小排序，大的可能是支付协议（多页），小的可能是结算单
        doc_images_sorted = sorted(doc_images, key=lambda x: x.stat().st_size, reverse=True)
        evidence_files['支付协议'].append(doc_images_sorted[0])  # 最大的
        evidence_files['结算单'].append(doc_images_sorted[1])   # 第二大的
    elif len(doc_images) == 1:
        evidence_files['支付协议'].append(doc_images[0])

    return evidence_files


def process_evidence_screenshots(project_path, evidence_files, screenshot_dir):
    """
    处理所有证据材料的截图
    - PDF: 提取关键页面
    - 图片: 复制到截图目录
    """
    screenshot_dir = Path(screenshot_dir)
    screenshot_dir.mkdir(exist_ok=True)

    all_screenshots = {}

    for evidence_type, files in evidence_files.items():
        if not files:
            continue

        type_screenshots = []

        for i, file_path in enumerate(files):
            file_path = Path(file_path)

            if file_path.suffix.lower() == '.pdf':
                # 提取PDF截图
                prefix = f"{evidence_type}_{i+1}"
                screenshots = extract_pdf_screenshots(file_path, screenshot_dir, prefix)
                type_screenshots.extend(screenshots)
            elif file_path.suffix.lower() in ['.png', '.jpg', '.jpeg']:
                # 复制图片到截图目录
                import shutil
                new_name = f"{evidence_type}_{i+1}{file_path.suffix}"
                dest_path = screenshot_dir / new_name
                shutil.copy(file_path, dest_path)
                type_screenshots.append(dest_path)

        all_screenshots[evidence_type] = type_screenshots

    return all_screenshots


def extract_pdf_screenshots(pdf_path, output_dir, prefix):
    """提取PDF关键页面截图"""
    try:
        import fitz
        screenshots = []

        doc = fitz.open(str(pdf_path))

        # 提取前3页（如果有）
        for page_num in range(min(3, len(doc))):
            page = doc[page_num]

            # 提高分辨率
            mat = fitz.Matrix(2, 2)
            pix = page.get_pixmap(matrix=mat)

            output_path = output_dir / f"{prefix}_page_{page_num + 1}.png"
            pix.save(str(output_path))
            screenshots.append(output_path)

        doc.close()
        return screenshots
    except Exception as e:
        print(f"[警告] PDF截图提取失败: {e}")
        return []


def generate_detailed_analysis(case_data, material_name, material_type):
    """生成详细分析内容"""

    analyses = {
        '中标公告': {
            'beneficial': '证明项目背景、千业公司作为联合中标人的主体资格，以及项目基本信息（中标价约7855万元，工期180天）。该证据可证明被告是项目的合法承包方，具备支付工程款的能力和义务。',
            'unfavorable': '仅证明中标事实，不能直接证明分包合同关系。需结合其他证据证明原告实际参与工程施工。',
            'opponent': '对方可能质疑该公告与本案的关联性，主张原告并非中标联合体成员，无权向其主张工程款。',
            'summary': '该证据主要用于证明项目背景和被告主体资格，建议结合施工合同、结算单等证据形成完整证据链。'
        },
        '支付协议': {
            'beneficial': '核心证据。协议确认：（1）甲方确认应付工程款145,565元；（2）分期付款安排明确；（3）被告已部分履行（付款50,000元），构成对协议的追认。该协议是主张欠款的最直接证据。',
            'unfavorable': '（1）协议由"翟振"签字，需证明其有权代表被告；（2）签约主体为"项目部"，存在主体瑕疵；（3）约定"剩余部分协商处理"，可能被对方利用抗辩。',
            'opponent': '（1）否认翟振的签字权限，主张协议无效；（2）以项目部无签约资格为由否认协议效力；（3）主张"协商处理"意味着未达成支付合意，拒绝支付余款。',
            'summary': '该证据是本案核心，但存在签字效力和主体瑕疵风险，需补强翟振身份证明和被告追认证据。'
        },
        '银行回单': {
            'beneficial': '付款凭证，证明被告已实际支付50,000元（2026年1月30日20,000元+2026年2月14日30,000元）。该证据可证明：（1）被告认可债务存在；（2）被告实际履行了部分付款义务，构成对协议的追认。',
            'unfavorable': '仅能证明已付款金额，不能证明欠款金额。需结合支付协议和结算单计算剩余欠款。',
            'opponent': '可能主张该付款是其他款项，与本案工程款无关；或主张已超额支付，要求返还。',
            'summary': '该证据是证明被告追认协议的关键，建议调取完整银行流水，固定付款事实。'
        },
        '结算单': {
            'beneficial': '工程款计算依据详细明确：（1）防水施工3679㎡×35元/㎡=128,765元；（2）点工47工日×400元/工日=18,800元；（3）扣减垃圾清理费2,000元；（4）合计145,565元。该结算单与支付协议金额一致，可互相印证。',
            'unfavorable': '（1）需核实结算单签字人身份和权限；（2）如被告否认结算单效力，需补充实际施工量证据；（3）点工部分缺乏详细工作内容记录。',
            'opponent': '（1）否认结算单签字人权限，主张结算无效；（2）质疑工程量和单价计算，要求重新核算；（3）主张存在质量问题应扣减工程款。',
            'summary': '该证据是计算工程款的重要依据，建议补充工程量确认单、验收记录等证据补强。'
        }
    }

    for key, value in analyses.items():
        if key in material_name:
            return value

    return {
        'beneficial': f'该材料与{case_data.get("案由", "本案")}相关，可作为证据支持我方主张。',
        'unfavorable': '需核实该证据的真实性、合法性、关联性。',
        'opponent': f'对方可能质疑该证据的真实性或关联性。',
        'summary': '该证据需与其他材料结合使用。'
    }


def generate_markdown_report(project_path, case_data, evidence_files, all_screenshots, output_md_path):
    """
    生成Markdown格式的材料分析报告
    """
    plaintiff = case_data.get('原告公司名称', case_data.get('原告姓名', '原告'))
    defendant = case_data.get('被告公司名称', case_data.get('被告姓名', '被告'))
    total_files = sum(len(files) for files in evidence_files.values())

    timestamp = datetime.now().strftime('%Y年%m月%d日')

    md_content = f"""# 案件材料分析报告

**{plaintiff} vs {defendant}**

---

## 一、案件基本信息

| 项目 | 内容 |
|------|------|
| 案由 | {case_data.get('案由', '待补充')} |
| 委托人（原告） | {plaintiff} |
| 对方当事人（被告） | {defendant} |
| 受诉法院 | {case_data.get('受诉法院', '待补充')} |
| 案号 | {case_data.get('案号', '待补充')} |
| 承办律师 | 待补充 |
| 分析日期 | {timestamp} |

---

## 二、工程项目背景

根据现有材料，本案涉及的工程项目信息如下：

**合同关系链**：{plaintiff}（防水工程分包方）→ {defendant}（净化工程总包方之一）→ 龙港市人民医院项目（业主）。根据中标结果公告，项目中标价约7855万元，工期180天。

**工程内容**：龙港市人民医院净化工程防水施工，包括防水施工3679㎡、点工47工日等。

**合同履行情况**：双方于2025年12月22日签订《支付协议》，确认工程款145,565元，约定分期支付。被告已支付50,000元，剩余95,565元未付。

---

## 三、证据材料清单与分析

### （一）材料清单

本次共扫描到{total_files}份证据材料：

"""

    # 列出证据文件
    evidence_num = 1
    for etype, files in evidence_files.items():
        if files:
            for f in files:
                md_content += f"{evidence_num}. {f.name}（{f.suffix[1:].upper()}，{f.stat().st_size / 1024:.1f} KB）\n"
                evidence_num += 1

    md_content += "\n### （二）逐份证据详细分析\n\n"

    # 证据1：中标公告
    if evidence_files['中标公告']:
        analysis = generate_detailed_analysis(case_data, '中标公告', '中标公告')
        md_content += f"""#### 【证据1：中标结果公告】

**1. 有利点分析**

{analysis['beneficial']}

**2. 关键内容截图**

"""
        for i, screenshot in enumerate(all_screenshots.get('中标公告', []), 1):
            md_content += f"![图1-{i}：中标结果公告关键内容](../截图证据/{screenshot.name})\n\n"

        md_content += f"""**3. 不利点及风险**

{analysis['unfavorable']}

**4. 相对方可能攻击点**

{analysis['opponent']}

**5. 分析小结**

{analysis['summary']}

---

"""

    # 证据2：支付协议
    if evidence_files['支付协议']:
        analysis = generate_detailed_analysis(case_data, '支付协议', '支付协议')
        md_content += f"""#### 【证据2：支付协议（2025年12月22日）】

**1. 有利点分析**

{analysis['beneficial']}

**2. 关键内容截图**

"""
        for i, screenshot in enumerate(all_screenshots.get('支付协议', []), 1):
            md_content += f"![图2-{i}：支付协议关键内容](../截图证据/{screenshot.name})\n\n"

        md_content += f"""**3. 不利点及风险**

{analysis['unfavorable']}

**4. 相对方可能攻击点**

{analysis['opponent']}

**5. 分析小结**

{analysis['summary']}

---

"""

    # 证据3：银行回单
    if evidence_files['银行回单']:
        analysis = generate_detailed_analysis(case_data, '银行回单', '银行回单')
        md_content += f"""#### 【证据3：银行电子回单×2】

**1. 有利点分析**

{analysis['beneficial']}

**2. 关键内容截图**

"""
        for i, screenshot in enumerate(all_screenshots.get('银行回单', []), 1):
            md_content += f"![图3-{i}：银行电子回单关键内容](../截图证据/{screenshot.name})\n\n"

        md_content += f"""**3. 不利点及风险**

{analysis['unfavorable']}

**4. 相对方可能攻击点**

{analysis['opponent']}

**5. 分析小结**

{analysis['summary']}

---

"""

    # 证据4：结算单
    if evidence_files['结算单']:
        analysis = generate_detailed_analysis(case_data, '结算单', '结算单')
        md_content += f"""#### 【证据4：劳务分包结算单（2024年12月27日）】

**1. 有利点分析**

{analysis['beneficial']}

**2. 关键内容截图**

"""
        for i, screenshot in enumerate(all_screenshots.get('结算单', []), 1):
            md_content += f"![图4-{i}：结算单关键内容](../截图证据/{screenshot.name})\n\n"

        md_content += f"""**3. 不利点及风险**

{analysis['unfavorable']}

**4. 相对方可能攻击点**

{analysis['opponent']}

**5. 分析小结**

{analysis['summary']}

---

"""

    # 欠款金额计算
    md_content += """### （三）欠款金额计算

1. **总工程款**：145,565元（根据支付协议和结算单确认）
2. **已付款**：50,000元（银行电子回单×2证明）
   - 2026年1月30日付款：20,000元
   - 2026年2月14日付款：30,000元
3. **剩余未付款**：95,565元（计算：145,565 - 50,000 = 95,565）
4. **逾期利息**：以95,565元为基数，自2026年2月17日起按LPR计算至实际付清之日止

---

## 四、核心争议焦点

| 风险等级 | 争议焦点 | 说明 |
|---------|---------|------|
| 🔴 高风险 | 项目经理签字效力问题 | 支付协议由乙方经办人"翟振"签字，但未见授权委托书或职务证明。{defendant}后续付款行为可视为对协议的追认，但仍需补强翟振的身份及权限证明。 |
| 🟡 中风险 | 协议主体瑕疵问题 | 签约乙方为"龙港市人民医院净化工程项目部"，该主体并非独立法人。虽然{defendant}实际履行了部分付款义务，但对方仍可能以此抗辩协议效力。 |
| 🟡 中风险 | 剩余款项支付条件争议 | 协议第二条约定"剩余部分甲乙双方协商处理"，该表述模糊，对方可能主张剩余款项支付条件尚未成就。 |

---

## 五、风险提示

### 🔴 高风险：项目经理签字效力

如无法证明翟振具有代表千业公司签署付款协议的权限，对方可能否认协议效力。建议立即收集：翟振的授权委托书、职务证明、社保缴纳记录等。

### 🟡 中风险：协议主体瑕疵

签约主体为项目部，存在主体不适格风险。{defendant}后续付款行为可作为追认证据。

### 🟡 中风险：剩余款项抗辩

协议"剩余部分协商处理"的表述可能被对方利用。建议收集协商过程的证据（微信记录、通话录音等）。

### 🟢 低风险：诉讼时效

第二期付款期限为2026年2月16日，目前仍在诉讼时效内。但建议尽快采取法律行动。

---

## 六、证据缺口与补强建议

| 序号 | 证据缺口 | 补强建议 |
|------|---------|---------|
| 1 | 建设工程施工合同 | 需收集原件，证明合同关系及工程内容 |
| 2 | 项目经理身份证明及授权文件 | **关键**：证明"翟振"签字效力 |
| 3 | 工程验收/结算文件 | 除结算单外，需收集竣工验收文件等 |
| 4 | 催款记录 | 证明时效及对方违约 |
| 5 | 损失计算依据 | 逾期利息、违约金等计算依据 |

---

## 七、证据链分析

现有{total_files}份证据材料已形成初步证据链：

1. **【中标结果公告】** → 证明项目背景和被告主体资格
2. **【支付协议】** → 确认债务金额和付款安排（核心证据）
3. **【银行电子回单】** → 证明被告部分履行，构成追认
4. **【劳务分包结算单】** → 印证工程款计算依据

**证据链完整性评估**：现有证据已能证明债务存在和欠款金额，但签字效力和主体瑕疵问题需要补强。

---

## 八、综合结论

经对{plaintiff}与{defendant}建设工程合同纠纷案件现有材料进行详细分析，共梳理出{total_files}份关键证据材料。已确认总工程款145,565元，已付款50,000元，剩余95,565元未付。

### 优势分析

1. 支付协议明确了债务金额，是最直接的债权凭证
2. 被告已部分履行（付款50,000元），构成对协议的追认
3. 银行回单和结算单可互相印证，证据链较为完整

### 风险提示

1. 项目经理签字效力和协议主体瑕疵是主要风险点
2. "剩余部分协商处理"的表述可能被对方利用抗辩

### 建议方案

建议尽快补强关键证据后，及时向龙港市人民法院提起诉讼。

---

**报告生成时间**：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
"""

    # 保存Markdown文件
    with open(output_md_path, 'w', encoding='utf-8') as f:
        f.write(md_content)

    print(f"[✓] Markdown报告已生成: {output_md_path}")
    return output_md_path


def convert_markdown_to_word(md_path, output_docx_path):
    """
    使用pandoc将Markdown转换为Word文档
    """
    try:
        # 使用pandoc转换
        result = subprocess.run(
            ['pandoc', str(md_path), '-o', str(output_docx_path),
             '--reference-doc', '~/templates/参考模板.docx'],
            capture_output=True,
            text=True
        )

        if result.returncode == 0:
            print(f"[✓] Word文档已生成: {output_docx_path}")
            return True
        else:
            # 如果pandoc失败，使用备用方案（python-docx）
            print(f"[!] pandoc转换失败，使用备用方案...")
            return convert_with_python_docx(md_path, output_docx_path)

    except FileNotFoundError:
        # pandoc未安装，使用备用方案
        print(f"[!] pandoc未安装，使用备用方案...")
        return convert_with_python_docx(md_path, output_docx_path)


def convert_with_python_docx(md_path, output_docx_path):
    """
    备用方案：使用python-docx从Markdown生成Word
    """
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    def set_chinese_font(run, font_name='宋体', font_size=10.5, bold=False, color=None):
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
        if color:
            run.font.color.rgb = color
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    doc = Document()

    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # 解析Markdown并转换为Word
    lines = md_content.split('\n')
    for line in lines:
        if line.startswith('# '):
            # 标题
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            set_chinese_font(run, '黑体', 22, bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            # 一级标题
            p = doc.add_paragraph()
            run = p.add_run(line[3:])
            set_chinese_font(run, '黑体', 16, bold=True)
        elif line.startswith('### '):
            # 二级标题
            p = doc.add_paragraph()
            run = p.add_run(line[4:])
            set_chinese_font(run, '黑体', 14, bold=True)
        elif line.startswith('#### '):
            # 三级标题
            p = doc.add_paragraph()
            run = p.add_run(line[5:])
            set_chinese_font(run, '黑体', 10.5, bold=True)
        elif line.startswith('**') and line.endswith('**'):
            # 加粗文本
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            set_chinese_font(run, '黑体', 10.5, bold=True)
        elif line.startswith('---'):
            # 分隔线
            doc.add_paragraph()
        elif line.startswith('![图'):
            # 图片引用（需要在后续处理中嵌入实际图片）
            p = doc.add_paragraph()
            run = p.add_run(line)
            set_chinese_font(run, '宋体', 10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('|'):
            # 表格行（简化处理）
            p = doc.add_paragraph()
            run = p.add_run(line)
            set_chinese_font(run, '宋体', 10.5)
        elif line.strip():
            # 普通文本
            p = doc.add_paragraph()
            run = p.add_run(line)
            set_chinese_font(run, '宋体', 10.5)

    doc.save(str(output_docx_path))
    print(f"[✓] Word文档已生成（备用方案）: {output_docx_path}")
    return True


def generate_comprehensive_report(project_path, case_data, output_path=None):
    """
    生成完整详细的材料分析报告（Markdown优先版）
    流程：Markdown → Word
    """
    project_path = Path(project_path)

    # 创建markdown目录
    markdown_dir = project_path / 'markdown'
    markdown_dir.mkdir(exist_ok=True)

    # 扫描所有证据文件
    print("[*] 扫描证据文件...")
    evidence_files = scan_evidence_files(project_path)
    print(f"[*] 发现证据：")
    for etype, files in evidence_files.items():
        if files:
            print(f"    - {etype}: {len(files)}个文件")

    # 处理截图
    print("[*] 提取证据截图...")
    screenshot_dir = project_path / '截图证据'
    all_screenshots = process_evidence_screenshots(project_path, evidence_files, screenshot_dir)

    # 生成时间戳
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

    # 第一步：生成Markdown报告
    md_output_path = markdown_dir / f'材料分析报告_{timestamp}.md'
    generate_markdown_report(project_path, case_data, evidence_files, all_screenshots, md_output_path)

    # 第二步：转换为Word文档
    if output_path is None:
        output_path = project_path / f'材料分析报告_{timestamp}.docx'

    convert_markdown_to_word(md_output_path, output_path)

    print(f"\n[*] 完整分析报告生成完成！")
    print(f"[*] Markdown报告: {md_output_path}")
    print(f"[*] Word文档: {output_path}")

    return str(output_path)


if __name__ == '__main__':
    import argparse

    parser = argparse.ArgumentParser(description='材料分析工具（v4.0 Markdown优先版）')
    parser.add_argument('--project', '-p', required=True, help='项目路径')
    parser.add_argument('--data', '-d', required=True, help='案件数据JSON文件路径')
    parser.add_argument('--output', '-o', help='输出报告路径（可选）')

    args = parser.parse_args()

    with open(args.data, 'r', encoding='utf-8') as f:
        case_data = json.load(f)

    report_path = generate_comprehensive_report(args.project, case_data, args.output)

    print(f"\n[*] 完整分析报告生成完成！")
    print(f"[*] 报告路径: {report_path}")