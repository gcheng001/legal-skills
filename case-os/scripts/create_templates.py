#!/usr/bin/env python3
"""
创建填空版模板
使用 {{字段名}} 占位符
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

def set_chinese_font(run, font_name='仿宋', font_size=14):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def create_complaint_template():
    """创建民事起诉状通用模板"""
    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # 标题
    title = doc.add_paragraph()
    title_run = title.add_run('民事起诉状')
    set_chinese_font(title_run, '黑体', 22)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(20)

    # 原告信息
    p = doc.add_paragraph()
    run = p.add_run('原告：')
    set_chinese_font(run, '黑体', 14)
    run = p.add_run('{{原告姓名}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('身份证号：')
    set_chinese_font(run, '仿宋', 14)
    run = p.add_run('{{原告身份证号}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('住址：')
    set_chinese_font(run, '仿宋', 14)
    run = p.add_run('{{原告地址}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('联系电话：')
    set_chinese_font(run, '仿宋', 14)
    run = p.add_run('{{原告电话}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()  # 空行

    # 被告信息
    p = doc.add_paragraph()
    run = p.add_run('被告：')
    set_chinese_font(run, '黑体', 14)
    run = p.add_run('{{被告姓名}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('身份证号：')
    set_chinese_font(run, '仿宋', 14)
    run = p.add_run('{{被告身份证号}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('住址：')
    set_chinese_font(run, '仿宋', 14)
    run = p.add_run('{{被告地址}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()  # 空行

    # 案由
    p = doc.add_paragraph()
    run = p.add_run('案由：')
    set_chinese_font(run, '黑体', 14)
    run = p.add_run('{{案由}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()  # 空行

    # 诉讼请求
    p = doc.add_paragraph()
    run = p.add_run('诉讼请求：')
    set_chinese_font(run, '黑体', 14)

    p = doc.add_paragraph()
    run = p.add_run('1、{{诉讼请求1}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('2、{{诉讼请求2}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('3、{{诉讼请求3}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()  # 空行

    # 事实与理由
    p = doc.add_paragraph()
    run = p.add_run('事实与理由：')
    set_chinese_font(run, '黑体', 14)

    p = doc.add_paragraph()
    run = p.add_run('{{事实与理由}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()  # 空行

    # 证据清单
    p = doc.add_paragraph()
    run = p.add_run('证据清单：')
    set_chinese_font(run, '黑体', 14)

    p = doc.add_paragraph()
    run = p.add_run('{{证据清单}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()  # 空行
    doc.add_paragraph()  # 空行

    # 此致
    p = doc.add_paragraph()
    run = p.add_run('此致')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('{{受诉法院}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()  # 空行
    doc.add_paragraph()  # 空行

    # 具状人和日期（右对齐）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('具状人(签字、盖章)：{{具状人}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('日期：{{日期}}')
    set_chinese_font(run, '仿宋', 14)

    return doc

def create_power_of_attorney_template():
    """创建委托手续模板"""
    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # 标题
    title = doc.add_paragraph()
    title_run = title.add_run('委托手续')
    set_chinese_font(title_run, '黑体', 22)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(20)

    # 委托书
    p = doc.add_paragraph()
    run = p.add_run('一、授权委托书')
    set_chinese_font(run, '黑体', 16)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('委托人：{{原告姓名}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('身份证号：{{原告身份证号}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('住址：{{原告地址}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('联系电话：{{原告电话}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('受委托人：{{代理人姓名}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('工作单位：{{律所名称}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('职务：律师')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('联系电话：{{代理人电话}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('现委托{{代理人姓名}}在{{原告姓名}}与{{被告姓名}}{{案由}}纠纷一案中，作为我方委托代理人。')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('委托事项和权限如下：')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('{{委托权限}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('委托人(签字)：{{原告姓名}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('日期：{{日期}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_page_break()

    # 所函
    p = doc.add_paragraph()
    run = p.add_run('二、律师事务所函')
    set_chinese_font(run, '黑体', 16)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('{{受诉法院}}：')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('本所接受{{原告姓名}}的委托，指派{{代理人姓名}}律师担任其与你院受理的{{案由}}纠纷一案中{{原告姓名}}的委托代理人。')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('特此函告')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('附：授权委托书一份')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('{{律所名称}}（盖章）')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('日期：{{日期}}')
    set_chinese_font(run, '仿宋', 14)

    return doc

def create_legal_representative_template():
    """创建法定代表人身份证明模板"""
    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # 标题
    title = doc.add_paragraph()
    title_run = title.add_run('法定代表人身份证明书')
    set_chinese_font(title_run, '黑体', 22)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(20)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('{{原告姓名}}（单位名称）在我单位任{{法定代表人职务}}职务，是我单位的法定代表人，特此证明。')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('附：法定代表人身份证复印件')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('单位（盖章）：{{原告姓名}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('日期：{{日期}}')
    set_chinese_font(run, '仿宋', 14)

    return doc

if __name__ == '__main__':
    import os

    base_path = '~/Codex/文件中转站/模板/民事案件'

    # 创建起诉状模板
    complaint_doc = create_complaint_template()
    complaint_path = os.path.join(base_path, '起诉状', '民事起诉状-通用模板.docx')
    complaint_doc.save(complaint_path)
    print(f'已创建: {complaint_path}')

    # 创建委托手续模板
    poa_doc = create_power_of_attorney_template()
    poa_path = os.path.join(base_path, '委托手续-填空模板.docx')
    poa_doc.save(poa_path)
    print(f'已创建: {poa_path}')

    # 创建法定代表人身份证明模板
    lr_doc = create_legal_representative_template()
    lr_path = os.path.join(base_path, '法定代表人身份证明-填空模板.docx')
    lr_doc.save(lr_path)
    print(f'已创建: {lr_path}')

    print('\n三个填空版模板已创建完成！')
    print('占位符格式: {{字段名}}')
