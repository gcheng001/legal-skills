#!/usr/bin/env python3
"""
立案材料生成脚本
根据案件信息自动填空模板，生成完整的立案材料包
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import json
import os
import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

# 答辩状案由映射（被告专用）
DEFENSE_CASE_TYPES = {
    '建设工程施工合同': '建设工程施工合同纠纷答辩状.docx',
    '建设工程': '建设工程施工合同纠纷答辩状.docx',
    '劳动': '劳动争议答辩状.docx',
    '劳务': '劳动争议答辩状.docx',
    '劳动合同': '劳动争议答辩状.docx',
    '民间借贷': '民间借贷纠纷答辩状.docx',
    '借款': '民间借贷纠纷答辩状.docx',
    '买卖': '买卖合同纠纷答辩状.docx',
    '购销': '买卖合同纠纷答辩状.docx',
    '机动车交通事故': '机动车交通事故责任纠纷答辩状.docx',
    '交通事故': '机动车交通事故责任纠纷答辩状.docx',
    '离婚': '离婚纠纷答辩状.docx',
    '婚姻': '离婚纠纷答辩状.docx',
    '租赁': '房屋租赁合同纠纷答辩状.docx',
    '房屋租赁': '房屋租赁合同纠纷答辩状.docx',
    '服务': '服务合同纠纷答辩状.docx',
    '股权转让': '股权转让纠纷答辩状.docx',
    '继承': '继承纠纷答辩状.docx',
    '侵权': '侵权责任纠纷答辩状.docx',
}

# 要素式起诉状案由映射
ELEMENTARY_CASE_TYPES = {
    '建设工程施工合同': '建设工程施工合同纠纷起诉状.docx',
    '建设工程': '建设工程施工合同纠纷起诉状.docx',
    '劳动': '劳动争议起诉状.docx',
    '劳务': '劳动争议起诉状.docx',
    '劳动合同': '劳动争议起诉状.docx',
    '民间借贷': '民间借贷纠纷起诉状.docx',
    '借款': '民间借贷纠纷起诉状.docx',
    '买卖': '买卖合同纠纷起诉状.docx',
    '购销': '买卖合同纠纷起诉状.docx',
    '机动车交通事故': '机动车交通事故责任纠纷起诉状.docx',
    '交通事故': '机动车交通事故责任纠纷起诉状.docx',
    '车祸': '机动车交通事故责任纠纷起诉状.docx',
    '离婚': '离婚纠纷起诉状.docx',
    '婚姻': '离婚纠纷起诉状.docx',
    '保证保险': '保证保险合同纠纷起诉状.docx',
    '金融借款': '金融借款合同纠纷起诉状.docx',
    '贷款': '金融借款合同纠纷起诉状.docx',
    '融资租赁': '融资租赁合同纠纷起诉状.docx',
    '物业': '物业服务合同纠纷起诉状.docx',
    '物业服务': '物业服务合同纠纷起诉状.docx',
    '信用卡': '银行信用卡纠纷起诉状.docx',
    '银行卡': '银行信用卡纠纷起诉状.docx',
    '证券': '证券虚假陈述责任纠纷起诉状.docx',
    '虚假陈述': '证券虚假陈述责任纠纷起诉状.docx',
}

def set_chinese_font(run, font_name='仿宋', font_size=14):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def extract_placeholders(doc):
    """提取文档中的所有占位符"""
    placeholders = set()
    placeholder_pattern = r'\{\{([^}]+)\}\}'

    for para in doc.paragraphs:
        matches = re.findall(placeholder_pattern, para.text)
        placeholders.update(matches)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    matches = re.findall(placeholder_pattern, para.text)
                    placeholders.update(matches)

    return placeholders

def replace_placeholders_in_paragraph(para, data):
    """替换段落中的占位符，保持格式"""
    placeholder_pattern = r'\{\{([^}]+)\}\}'

    # 获取段落完整文本
    full_text = para.text

    # 查找所有占位符
    matches = list(re.finditer(placeholder_pattern, full_text))
    if not matches:
        return False

    # 如果有匹配，需要重建段落
    if matches:
        # 保存第一个run的格式
        first_run = para.runs[0] if para.runs else None
        if first_run:
            font_name = first_run.font.name or '仿宋'
            font_size = first_run.font.size.pt if first_run.font.size else 14
            bold = first_run.bold
        else:
            font_name = '仿宋'
            font_size = 14
            bold = False

        # 清除所有runs
        for run in para.runs:
            run._element.getparent().remove(run._element)

        # 重建文本
        last_end = 0
        for match in matches:
            placeholder = match.group(0)  # {{字段名}}
            field_name = match.group(1)   # 字段名
            start, end = match.span()

            # 添加占位符前的文本
            if start > last_end:
                text_before = full_text[last_end:start]
                run = para.add_run(text_before)
                set_chinese_font(run, font_name, font_size)
                if bold:
                    run.bold = True

            # 添加替换后的值
            replacement = data.get(field_name, placeholder)
            run = para.add_run(replacement)
            set_chinese_font(run, font_name, font_size)
            if bold:
                run.bold = True

            last_end = end

        # 添加剩余文本
        if last_end < len(full_text):
            text_after = full_text[last_end:]
            run = para.add_run(text_after)
            set_chinese_font(run, font_name, font_size)
            if bold:
                run.bold = True

        return True

    return False

def fill_template(template_path, output_path, data):
    """填充模板"""
    doc = Document(template_path)

    # 处理段落
    for para in doc.paragraphs:
        replace_placeholders_in_paragraph(para, data)

    # 处理表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_placeholders_in_paragraph(para, data)

    # 保存
    doc.save(output_path)
    return True

def detect_case_type(case_overview, role='plaintiff'):
    """从案件概览中检测案由类型"""
    case_cause = case_overview.get('案由', '')

    if role == 'defendant':
        for keyword, template_file in DEFENSE_CASE_TYPES.items():
            if keyword in case_cause:
                return template_file
        return None

    for keyword, template_file in ELEMENTARY_CASE_TYPES.items():
        if keyword in case_cause:
            return template_file

    return None

def select_address_confirmation(court_name):
    """选择地址确认书版本"""
    if '鹿城区' in court_name or '鹿城' in court_name:
        return '地址确认书-鹿城区人民法院.docx'
    else:
        return '送达地址确认书.doc'

def check_final_gate(project_path, role='plaintiff'):
    """
    执行 FINAL 门禁检查
    注意：无论原告还是被告，均读取案件级原告链 S10 门禁结果
    （因为 S10 仅在原告九步法链生成，被告复用同一阻断逻辑）
    Returns: (passed, message)
    """
    import sys
    script_dir = Path(__file__).parent
    gate_script = script_dir / "check_final_gate.py"
    if not gate_script.exists():
        return False, "门禁脚本不存在"

    import subprocess
    # 统一使用 plaintiff 链 S10 文件做门禁判断
    result = subprocess.run(
        [sys.executable, str(gate_script), '--project', project_path, '--role', 'plaintiff'],
        capture_output=True, text=True
    )
    if result.returncode == 0:
        return True, result.stdout.strip()
    else:
        return False, result.stdout.strip() or result.stderr.strip()


def generate_filing_materials(project_path, case_data):
    """
    生成立案材料

    Args:
        project_path: 项目路径
        case_data: 案件数据字典

    Returns:
        dict: 生成的文件清单
    """
    # 门禁检查
    gate_passed, gate_msg = check_final_gate(project_path, 'plaintiff')
    if not gate_passed:
        print(f"[✗] 门禁未通过，无法生成立案材料")
        print(f"[✗] 原因: {gate_msg}")
        return {
            'success': False,
            'blocked': True,
            'block_reason': gate_msg
        }
    print(f"[*] 门禁检查通过: {gate_msg}")

    # 模板基础路径
    template_base = Path('~/templates/civil')

    # 创建立案材料文件夹
    output_dir = Path(project_path) / '立案材料-Word'
    output_dir.mkdir(exist_ok=True)

    generated_files = []

    # 准备填空数据
    fill_data = {
        # 原告信息
        '原告姓名': case_data.get('原告姓名', ''),
        '原告性别': case_data.get('原告性别', ''),
        '原告民族': case_data.get('原告民族', ''),
        '原告出生日期': case_data.get('原告出生日期', ''),
        '原告身份证号': case_data.get('原告身份证号', ''),
        '原告地址': case_data.get('原告地址', ''),
        '原告电话': case_data.get('原告电话', ''),
        '原告公司名称': case_data.get('原告公司名称', ''),
        '原告公司代码': case_data.get('原告公司代码', ''),
        '原告公司地址': case_data.get('原告公司地址', ''),
        '原告法定代表人': case_data.get('原告法定代表人', ''),
        '原告法定代表人职务': case_data.get('原告法定代表人职务', ''),
        '原告公司电话': case_data.get('原告公司电话', ''),

        # 被告信息
        '被告姓名': case_data.get('被告姓名', ''),
        '被告性别': case_data.get('被告性别', ''),
        '被告民族': case_data.get('被告民族', ''),
        '被告出生日期': case_data.get('被告出生日期', ''),
        '被告身份证号': case_data.get('被告身份证号', ''),
        '被告地址': case_data.get('被告地址', ''),
        '被告电话': case_data.get('被告电话', ''),
        '被告公司名称': case_data.get('被告公司名称', ''),
        '被告公司代码': case_data.get('被告公司代码', ''),
        '被告公司地址': case_data.get('被告公司地址', ''),
        '被告法定代表人': case_data.get('被告法定代表人', ''),
        '被告法定代表人职务': case_data.get('被告法定代表人职务', ''),
        '被告公司电话': case_data.get('被告公司电话', ''),

        # 案件信息
        '案由': case_data.get('案由', ''),
        '案号': case_data.get('案号', ''),
        '诉讼请求1': case_data.get('诉讼请求1', ''),
        '诉讼请求2': case_data.get('诉讼请求2', ''),
        '诉讼请求3': case_data.get('诉讼请求3', ''),
        '诉讼请求4': case_data.get('诉讼请求4', ''),
        '诉讼请求5': case_data.get('诉讼请求5', ''),
        '事实与理由': case_data.get('事实与理由', ''),
        '受诉法院': case_data.get('受诉法院', ''),
        '具状人': case_data.get('原告姓名', ''),
        '日期': datetime.now().strftime('%Y年%m月%d日'),

        # 委托手续信息
        '委托人姓名': case_data.get('原告姓名', ''),
        '委托人身份证号': case_data.get('原告身份证号', case_data.get('原告公司代码', '')),
        '委托人地址': case_data.get('原告地址', case_data.get('原告公司地址', '')),
        '委托人电话': case_data.get('原告电话', case_data.get('原告公司电话', '')),
        '对方姓名': case_data.get('被告姓名', case_data.get('被告公司名称', '')),
        '收费金额': case_data.get('收费金额', ''),
        '收费方式': case_data.get('收费方式', ''),
        '付款期限': case_data.get('付款期限', ''),
        '委托权限': case_data.get('委托权限', ''),
        '签订日期': datetime.now().strftime('%Y年%m月%d日'),
        '委托人签字': case_data.get('原告姓名', ''),
        '有效期限': case_data.get('有效期限', '30'),

        # 法定代表人信息
        '法定代表人姓名': case_data.get('原告法定代表人', ''),
        '法定代表人职务': case_data.get('原告法定代表人职务', ''),

        # 证据清单信息
        '提交时间': datetime.now().strftime('%Y年%m月%d日'),
        '提交人签字': case_data.get('原告姓名', ''),
    }

    # 添加证据信息（如果有）
    for i in range(1, 6):
        fill_data[f'证据{i}名称'] = case_data.get(f'证据{i}名称', '')
        fill_data[f'证据{i}来源'] = case_data.get(f'证据{i}来源', '')
        fill_data[f'证据{i}证明内容'] = case_data.get(f'证据{i}证明内容', '')
        fill_data[f'证据{i}页码'] = case_data.get(f'证据{i}页码', '')

    print(f"[*] 正在生成立案材料...")
    print(f"[*] 输出目录: {output_dir}")

    # 1. 生成普通版起诉状
    try:
        template_path = template_base / '起诉状' / '民事起诉状-通用模板.docx'
        output_path = output_dir / '起诉状.docx'
        fill_template(str(template_path), str(output_path), fill_data)
        generated_files.append('起诉状.docx')
        print(f"[✓] 已生成: 起诉状.docx")
    except Exception as e:
        print(f"[✗] 生成起诉状失败: {e}")

    # 2. 判断是否需要素式起诉状
    elementary_template = detect_case_type(case_data)
    if elementary_template:
        try:
            template_path = template_base / '起诉状' / '要素式起诉状' / elementary_template
            if template_path.exists():
                output_path = output_dir / '起诉状-要素式.docx'
                fill_template(str(template_path), str(output_path), fill_data)
                generated_files.append('起诉状-要素式.docx')
                print(f"[✓] 已生成: 起诉状-要素式.docx (案由: {case_data.get('案由', '')})")
        except Exception as e:
            print(f"[✗] 生成要素式起诉状失败: {e}")

    # 3. 生成委托手续
    try:
        template_path = template_base / '委托手续-填空模板.docx'
        output_path = output_dir / '委托手续.docx'
        fill_template(str(template_path), str(output_path), fill_data)
        generated_files.append('委托手续.docx')
        print(f"[✓] 已生成: 委托手续.docx")
    except Exception as e:
        print(f"[✗] 生成委托手续失败: {e}")

    # 4. 生成地址确认书
    try:
        court_name = case_data.get('受诉法院', '')
        address_template = select_address_confirmation(court_name)

        if '鹿城区' in court_name:
            template_path = template_base / '地址确认书' / '地址确认书-鹿城区人民法院.docx'
        else:
            # 注意：送达地址确认书是.doc格式，需要转换或处理
            template_path = template_base / '地址确认书' / '送达地址确认书.doc'

        if template_path.exists():
            output_path = output_dir / '地址确认书.docx'
            if template_path.suffix == '.docx':
                fill_template(str(template_path), str(output_path), fill_data)
            else:
                # 对于.doc文件，先复制，后续手动处理
                shutil.copy(str(template_path), str(output_dir / '地址确认书.doc'))
                print(f"[!] 地址确认书为.doc格式，已复制到文件夹，请手动填写")
            generated_files.append('地址确认书.docx')
            print(f"[✓] 已生成: 地址确认书 ({'鹿城区版本' if '鹿城区' in court_name else '通用版本'})")
    except Exception as e:
        print(f"[✗] 生成地址确认书失败: {e}")

    # 5. 生成法定代表人身份证明（如果委托人为公司）
    if case_data.get('原告类型') == '公司' or case_data.get('原告公司名称'):
        try:
            template_path = template_base / '法定代表人身份证明-填空模板.docx'
            output_path = output_dir / '法定代表人身份证明.docx'
            fill_template(str(template_path), str(output_path), fill_data)
            generated_files.append('法定代表人身份证明.docx')
            print(f"[✓] 已生成: 法定代表人身份证明.docx")
        except Exception as e:
            print(f"[✗] 生成法定代表人身份证明失败: {e}")

    # 6. 生成证据清单
    try:
        template_path = template_base / '证据清单-填空模板.docx'
        output_path = output_dir / '证据清单.docx'
        fill_template(str(template_path), str(output_path), fill_data)
        generated_files.append('证据清单.docx')
        print(f"[✓] 已生成: 证据清单.docx")
    except Exception as e:
        print(f"[✗] 生成证据清单失败: {e}")

    print(f"\n[*] 立案材料生成完成！")
    print(f"[*] 共生成 {len(generated_files)} 个文件")
    print(f"[*] 输出目录: {output_dir}")

    return {
        'success': True,
        'output_dir': str(output_dir),
        'generated_files': generated_files,
        'count': len(generated_files)
    }

def generate_defense_materials(project_path, case_data):
    """
    生成被告答辩材料

    Args:
        project_path: 项目路径
        case_data: 案件数据字典

    Returns:
        dict: 生成的文件清单
    """
    # 正式答辩材料同样受案件级原告主链 S10 门禁控制。
    gate_passed, gate_msg = check_final_gate(project_path, 'plaintiff')
    if not gate_passed:
        print(f"[✗] 门禁未通过，无法生成答辩材料")
        print(f"[✗] 原因: {gate_msg}")
        return {
            'success': False,
            'blocked': True,
            'block_reason': gate_msg
        }
    print(f"[*] 门禁检查通过: {gate_msg}")

    template_base = Path('~/templates/civil')

    output_dir = Path(project_path) / '答辩材料-Word'
    output_dir.mkdir(exist_ok=True)

    generated_files = []

    # 被告视角填空数据（原告/被告字段互换语义）
    fill_data = {
        # 被告信息（我方当事人）
        '原告姓名': case_data.get('被告姓名', ''),
        '原告性别': case_data.get('被告性别', ''),
        '原告民族': case_data.get('被告民族', ''),
        '原告出生日期': case_data.get('被告出生日期', ''),
        '原告身份证号': case_data.get('被告身份证号', ''),
        '原告地址': case_data.get('被告地址', ''),
        '原告电话': case_data.get('被告电话', ''),
        '原告公司名称': case_data.get('被告公司名称', ''),
        '原告公司代码': case_data.get('被告公司代码', ''),
        '原告公司地址': case_data.get('被告公司地址', ''),
        '原告法定代表人': case_data.get('被告法定代表人', ''),
        '原告法定代表人职务': case_data.get('被告法定代表人职务', ''),
        '原告公司电话': case_data.get('被告公司电话', ''),

        # 原告信息（对方当事人）
        '被告姓名': case_data.get('原告姓名', ''),
        '被告性别': case_data.get('原告性别', ''),
        '被告民族': case_data.get('原告民族', ''),
        '被告出生日期': case_data.get('原告出生日期', ''),
        '被告身份证号': case_data.get('原告身份证号', ''),
        '被告地址': case_data.get('原告地址', ''),
        '被告电话': case_data.get('原告电话', ''),
        '被告公司名称': case_data.get('原告公司名称', ''),
        '被告公司代码': case_data.get('原告公司代码', ''),
        '被告公司地址': case_data.get('原告公司地址', ''),
        '被告法定代表人': case_data.get('原告法定代表人', ''),
        '被告法定代表人职务': case_data.get('原告法定代表人职务', ''),
        '被告公司电话': case_data.get('原告公司电话', ''),

        # 案件信息
        '案由': case_data.get('案由', ''),
        '案号': case_data.get('案号', ''),
        '答辩请求1': case_data.get('答辩请求1', '驳回原告全部诉讼请求'),
        '答辩请求2': case_data.get('答辩请求2', ''),
        '答辩请求3': case_data.get('答辩请求3', ''),
        '事实与理由': case_data.get('答辩意见', case_data.get('事实与理由', '')),
        '受诉法院': case_data.get('受诉法院', ''),
        '具状人': case_data.get('被告姓名', case_data.get('被告公司名称', '')),
        '日期': datetime.now().strftime('%Y年%m月%d日'),

        # 委托手续信息（被告方）
        '委托人姓名': case_data.get('被告姓名', case_data.get('被告公司名称', '')),
        '委托人身份证号': case_data.get('被告身份证号', case_data.get('被告公司代码', '')),
        '委托人地址': case_data.get('被告地址', case_data.get('被告公司地址', '')),
        '委托人电话': case_data.get('被告电话', case_data.get('被告公司电话', '')),
        '对方姓名': case_data.get('原告姓名', case_data.get('原告公司名称', '')),
        '收费金额': case_data.get('收费金额', ''),
        '收费方式': case_data.get('收费方式', ''),
        '付款期限': case_data.get('付款期限', ''),
        '委托权限': case_data.get('委托权限', ''),
        '签订日期': datetime.now().strftime('%Y年%m月%d日'),
        '委托人签字': case_data.get('被告姓名', case_data.get('被告公司名称', '')),
        '有效期限': case_data.get('有效期限', '30'),

        # 法定代表人信息（被告方）
        '法定代表人姓名': case_data.get('被告法定代表人', ''),
        '法定代表人职务': case_data.get('被告法定代表人职务', ''),

        # 证据清单信息
        '提交时间': datetime.now().strftime('%Y年%m月%d日'),
        '提交人签字': case_data.get('被告姓名', case_data.get('被告公司名称', '')),
    }

    # 添加证据信息
    for i in range(1, 6):
        fill_data[f'证据{i}名称'] = case_data.get(f'证据{i}名称', '')
        fill_data[f'证据{i}来源'] = case_data.get(f'证据{i}来源', '')
        fill_data[f'证据{i}证明内容'] = case_data.get(f'证据{i}证明内容', '')
        fill_data[f'证据{i}页码'] = case_data.get(f'证据{i}页码', '')

    print(f"[*] 正在生成被告答辩材料...")
    print(f"[*] 输出目录: {output_dir}")

    # 1. 生成答辩状
    try:
        template_path = template_base / '答辩状' / '民事答辩状-通用模板.docx'
        if template_path.exists():
            output_path = output_dir / '答辩状.docx'
            fill_template(str(template_path), str(output_path), fill_data)
            generated_files.append('答辩状.docx')
            print(f"[✓] 已生成: 答辩状.docx")
        else:
            print(f"[!] 答辩状模板不存在: {template_path}")
    except Exception as e:
        print(f"[✗] 生成答辩状失败: {e}")

    # 2. 判断是否需要要素式答辩状
    defense_template = detect_case_type(case_data, role='defendant')
    if defense_template:
        try:
            template_path = template_base / '答辩状' / '要素式答辩状' / defense_template
            if template_path.exists():
                output_path = output_dir / '答辩状-要素式.docx'
                fill_template(str(template_path), str(output_path), fill_data)
                generated_files.append('答辩状-要素式.docx')
                print(f"[✓] 已生成: 答辩状-要素式.docx (案由: {case_data.get('案由', '')})")
        except Exception as e:
            print(f"[✗] 生成要素式答辩状失败: {e}")

    # 3. 生成委托手续（被告方）
    try:
        template_path = template_base / '委托手续-填空模板.docx'
        output_path = output_dir / '委托手续-被告.docx'
        fill_template(str(template_path), str(output_path), fill_data)
        generated_files.append('委托手续-被告.docx')
        print(f"[✓] 已生成: 委托手续-被告.docx")
    except Exception as e:
        print(f"[✗] 生成委托手续失败: {e}")

    # 4. 生成地址确认书
    try:
        court_name = case_data.get('受诉法院', '')
        if '鹿城区' in court_name:
            template_path = template_base / '地址确认书' / '地址确认书-鹿城区人民法院.docx'
        else:
            template_path = template_base / '地址确认书' / '送达地址确认书.doc'

        if template_path.exists():
            output_path = output_dir / '地址确认书-被告.docx'
            if template_path.suffix == '.docx':
                fill_template(str(template_path), str(output_path), fill_data)
            else:
                shutil.copy(str(template_path), str(output_dir / '地址确认书-被告.doc'))
                print(f"[!] 地址确认书为.doc格式，已复制到文件夹，请手动填写")
            generated_files.append('地址确认书-被告.docx')
            print(f"[✓] 已生成: 地址确认书-被告")
    except Exception as e:
        print(f"[✗] 生成地址确认书失败: {e}")

    # 5. 生成法定代表人身份证明（被告方，如果被告为公司）
    if case_data.get('被告类型') == '公司' or case_data.get('被告公司名称'):
        try:
            template_path = template_base / '法定代表人身份证明-填空模板.docx'
            output_path = output_dir / '法定代表人身份证明-被告.docx'
            fill_template(str(template_path), str(output_path), fill_data)
            generated_files.append('法定代表人身份证明-被告.docx')
            print(f"[✓] 已生成: 法定代表人身份证明-被告.docx")
        except Exception as e:
            print(f"[✗] 生成法定代表人身份证明失败: {e}")

    # 6. 生成证据清单（被告方）
    try:
        template_path = template_base / '证据清单-填空模板.docx'
        output_path = output_dir / '证据清单-被告.docx'
        fill_template(str(template_path), str(output_path), fill_data)
        generated_files.append('证据清单-被告.docx')
        print(f"[✓] 已生成: 证据清单-被告.docx")
    except Exception as e:
        print(f"[✗] 生成证据清单失败: {e}")

    print(f"\n[*] 被告答辩材料生成完成！")
    print(f"[*] 共生成 {len(generated_files)} 个文件")
    print(f"[*] 输出目录: {output_dir}")

    return {
        'success': True,
        'role': 'defendant',
        'output_dir': str(output_dir),
        'generated_files': generated_files,
        'count': len(generated_files)
    }


def main():
    """主函数"""
    import argparse

    parser = argparse.ArgumentParser(description='生成立案材料/答辩材料')
    parser.add_argument('--project', required=True, help='项目路径')
    parser.add_argument('--data', required=True, help='案件数据JSON文件路径')
    parser.add_argument('--role', default='plaintiff', choices=['plaintiff', 'defendant'],
                        help='角色：plaintiff（原告，默认）或 defendant（被告）')
    args = parser.parse_args()

    # 读取案件数据
    with open(args.data, 'r', encoding='utf-8') as f:
        case_data = json.load(f)

    # 根据角色生成材料
    if args.role == 'defendant':
        result = generate_defense_materials(args.project, case_data)
    else:
        result = generate_filing_materials(args.project, case_data)

    # 输出结果
    print(json.dumps(result, ensure_ascii=False, indent=2))

    # 门禁阻断时返回非零退出码（不得在阻断前创建输出目录或文件）
    if result.get('blocked') or not result.get('success'):
        sys.exit(1)
    sys.exit(0)

if __name__ == '__main__':
    main()
