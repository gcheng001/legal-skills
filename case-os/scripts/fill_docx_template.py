#!/usr/bin/env python3
import argparse
import json
from datetime import datetime
from pathlib import Path

from docx import Document


def replace_in_runs(runs, replacements):
    changed = False
    for run in runs:
        text = run.text or ""
        new_text = text
        for old, new in replacements.items():
            if old and old in new_text:
                new_text = new_text.replace(old, new)
        if new_text != text:
            run.text = new_text
            changed = True
    return changed


def replace_everywhere(doc, replacements):
    changed = 0
    for p in doc.paragraphs:
        if replace_in_runs(p.runs, replacements):
            changed += 1
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if replace_in_runs(p.runs, replacements):
                        changed += 1
    return changed


def fill_element_complaint(doc, litigant_name, date_text):
    touched = 0
    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if txt.startswith("具状人"):
            p.text = f"具状人(签字、盖章): {litigant_name}"
            touched += 1
        elif txt.startswith("日期"):
            p.text = f"日期：{date_text}"
            touched += 1
    return touched


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--replacements", default="{}")
    parser.add_argument("--mode", default="replace", choices=["replace", "element_fill"])
    parser.add_argument("--litigant", default="")
    parser.add_argument("--date", default="")
    args = parser.parse_args()

    template = Path(args.template)
    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(template))
    replacements = json.loads(args.replacements or "{}")

    changed = replace_everywhere(doc, replacements)
    mode_touched = 0

    if args.mode == "element_fill":
        date_text = args.date or datetime.now().strftime("%Y年%m月%d日")
        mode_touched = fill_element_complaint(doc, args.litigant or "待补充", date_text)

    doc.save(str(output))

    print(
        json.dumps(
            {
                "ok": True,
                "template": str(template),
                "output": str(output),
                "changed_runs": changed,
                "mode_touched": mode_touched,
                "mode": args.mode,
            },
            ensure_ascii=False,
        )
    )


if __name__ == "__main__":
    main()
