#!/usr/bin/env python3
"""
材料分析脚本 - v3.0 智能截图版
自动识别所有证据材料（PDF+图片），提取/嵌入截图
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
import re
import json
from pathlib import Path
from datetime import datetime
from PIL import Image

# 导入PDF截图模块
import sys
sys.path.insert(0, str(Path(__file__).parent))
from extract_pdf_screenshot import extract_screenshots_for_evidence

def set_chinese_font(run, font_name='宋体', font_size=10.5, bold=False, color=None):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_title(doc, text, font_size=22):
    """添加标题（黑体加粗居中）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, '黑体', font_size, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(12)
    return p

def add_heading1(doc, text):
    """添加一级标题（黑体16号加粗）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, '黑体', 16, bold=True)
    p.space_after = Pt(8)
    return p

def add_heading2(doc, text):
    """添加二级标题（黑体14号加粗）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, '黑体', 14, bold=True)
    p.space_after = Pt(6)
    return p

def add_bold_label(doc, text):
    """添加加粗标签（黑体10.5号加粗）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, '黑体', 10.5, bold=True)
    return p

def add_normal_text(doc, text, bold=False):
    """添加正文（宋体10.5号）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, '宋体', 10.5, bold=bold)
    return p

def add_risk_section(doc, level, title, content):
    """添加风险提示（带颜色）"""
    p = doc.add_paragraph()

    if level == '高风险':
        level_run = p.add_run(f'{level}：')
        set_chinese_font(level_run, '黑体', 10.5, bold=True, color=RGBColor(255, 0, 0))
    elif level == '中风险':
        level_run = p.add_run(f'{level}：')
        set_chinese_font(level_run, '黑体', 10.5, bold=True, color=RGBColor(255, 165, 0))
    elif level == '低风险':
        level_run = p.add_run(f'{level}：')
        set_chinese_font(level_run, '黑体', 10.5, bold=True, color=RGBColor(0, 128, 0))
    else:
        level_run = p.add_run(f'{level}：')
        set_chinese_font(level_run, '黑体', 10.5, bold=True)

    title_run = p.add_run(title)
    set_chinese_font(title_run, '黑体', 10.5, bold=True)

    doc.add_paragraph()
    add_normal_text(doc, content)
    doc.add_paragraph()

def scan_evidence_files(project_path):
    """
    智能扫描所有证据文件
    返回按证据类型分类的文件列表
    """
    from PIL import Image

    project_path = Path(project_path)
    evidence_files = {
        '中标公告': [],
        '支付协议': [],
        '银行回单': [],
        '结算单': [],
        '其他': []
    }

    # 扫描所有文件
    all_files = []
    for file_path in project_path.rglob('*'):
        if file_path.is_file() and not any(x in str(file_path) for x in ['立案材料', '截图证据', '.temp_', '材料分析报告']):
            if file_path.suffix.lower() in ['.pdf', '.png', '.jpg', '.jpeg', '.doc', '.docx']:
                all_files.append(file_path)

    # 先按文件名分类
    unclassified = []
    for file_path in all_files:
        name_lower = file_path.name.lower()

        if '中标' in name_lower or '公告' in name_lower:
            evidence_files['中标公告'].append(file_path)
        elif '支付协议' in name_lower or '付款协议' in name_lower:
            evidence_files['支付协议'].append(file_path)
        elif '银行' in name_lower or '回单' in name_lower:
            evidence_files['银行回单'].append(file_path)
        elif '结算' in name_lower:
            evidence_files['结算单'].append(file_path)
        else:
            unclassified.append(file_path)

    # 智能分类未识别的图片（根据图片特征）
    bank_images = []
    doc_images = []

    for file_path in unclassified:
        if file_path.suffix.lower() in ['.png', '.jpg', '.jpeg']:
            try:
                img = Image.open(file_path)
                width, height = img.size
                ratio = width / height

                # 银行回单特征：横版，宽度>高度，比例>1.5
                if ratio > 1.5:
                    bank_images.append(file_path)
                # 文档照片特征：竖版，高度>宽度，比例<1.0
                elif ratio < 1.0:
                    doc_images.append(file_path)
                else:
                    evidence_files['其他'].append(file_path)
            except:
                evidence_files['其他'].append(file_path)
        else:
            evidence_files['其他'].append(file_path)

    # 分配银行回单（2个）
    if len(bank_images) >= 2:
        evidence_files['银行回单'] = bank_images[:2]

    # 分配文档（支付协议和结算单）
    if len(doc_images) >= 2:
        # 按文件大小排序，大的可能是支付协议（多页），小的可能是结算单
        doc_images_sorted = sorted(doc_images, key=lambda x: x.stat().st_size, reverse=True)
        evidence_files['支付协议'].append(doc_images_sorted[0])  # 最大的
        evidence_files['结算单'].append(doc_images_sorted[1])   # 第二大的
    elif len(doc_images) == 1:
        evidence_files['支付协议'].append(doc_images[0])

    return evidence_files

def embed_image_in_doc(doc, image_path, caption, width=Inches(5.5)):
    """在文档中嵌入图片"""
    try:
        # 打开图片获取尺寸
        img = Image.open(image_path)
        img_width, img_height = img.size

        # 计算缩放后的高度（保持比例）
        aspect_ratio = img_height / img_width
        new_height = width.inches * aspect_ratio

        # 添加图片
        doc.add_picture(str(image_path), width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加图注
        p = doc.add_paragraph()
        run = p.add_run(caption)
        set_chinese_font(run, '宋体', 10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        return True
    except Exception as e:
        add_normal_text(doc, f'[图片嵌入失败: {e}]')
        return False

def extract_pdf_screenshots(pdf_path, output_dir, prefix):
    """提取PDF关键页面截图"""
    try:
        import fitz
        screenshots = []

        doc = fitz.open(str(pdf_path))

        # 提取前3页（如果有）
        for page_num in range(min(3, len(doc))):
            page = doc[page_num]

            # 提高分辨率
            mat = fitz.Matrix(2, 2)
            pix = page.get_pixmap(matrix=mat)

            output_path = output_dir / f"{prefix}_page_{page_num + 1}.png"
            pix.save(str(output_path))
            screenshots.append(output_path)

        doc.close()
        return screenshots
    except Exception as e:
        print(f"[警告] PDF截图提取失败: {e}")
        return []

def process_evidence_screenshots(project_path, evidence_files, screenshot_dir):
    """
    处理所有证据材料的截图
    - PDF: 提取关键页面
    - 图片: 复制到截图目录
    """
    screenshot_dir = Path(screenshot_dir)
    screenshot_dir.mkdir(exist_ok=True)

    all_screenshots = {}

    for evidence_type, files in evidence_files.items():
        if not files:
            continue

        type_screenshots = []

        for i, file_path in enumerate(files):
            file_path = Path(file_path)

            if file_path.suffix.lower() == '.pdf':
                # 提取PDF截图
                prefix = f"{evidence_type}_{i+1}"
                screenshots = extract_pdf_screenshots(file_path, screenshot_dir, prefix)
                type_screenshots.extend(screenshots)
            elif file_path.suffix.lower() in ['.png', '.jpg', '.jpeg']:
                # 复制图片到截图目录
                import shutil
                new_name = f"{evidence_type}_{i+1}{file_path.suffix}"
                dest_path = screenshot_dir / new_name
                shutil.copy(file_path, dest_path)
                type_screenshots.append(dest_path)

        all_screenshots[evidence_type] = type_screenshots

    return all_screenshots

def generate_detailed_analysis(case_data, material_name, material_type):
    """生成详细分析内容"""

    analyses = {
        '中标公告': {
            'beneficial': '证明项目背景、千业公司作为联合中标人的主体资格，以及项目基本信息（中标价约7855万元，工期180天）。该证据可证明被告是项目的合法承包方，具备支付工程款的能力和义务。',
            'unfavorable': '仅证明中标事实，不能直接证明分包合同关系。需结合其他证据证明原告实际参与工程施工。',
            'opponent': '对方可能质疑该公告与本案的关联性，主张原告并非中标联合体成员，无权向其主张工程款。',
            'summary': '该证据主要用于证明项目背景和被告主体资格，建议结合施工合同、结算单等证据形成完整证据链。'
        },
        '支付协议': {
            'beneficial': '核心证据。协议确认：（1）甲方确认应付工程款145,565元；（2）分期付款安排明确；（3）被告已部分履行（付款50,000元），构成对协议的追认。该协议是主张欠款的最直接证据。',
            'unfavorable': '（1）协议由"翟振"签字，需证明其有权代表被告；（2）签约主体为"项目部"，存在主体瑕疵；（3）约定"剩余部分协商处理"，可能被对方利用抗辩。',
            'opponent': '（1）否认翟振的签字权限，主张协议无效；（2）以项目部无签约资格为由否认协议效力；（3）主张"协商处理"意味着未达成支付合意，拒绝支付余款。',
            'summary': '该证据是本案核心，但存在签字效力和主体瑕疵风险，需补强翟振身份证明和被告追认证据。'
        },
        '银行回单': {
            'beneficial': '付款凭证，证明被告已实际支付50,000元（2026年1月30日20,000元+2026年2月14日30,000元）。该证据可证明：（1）被告认可债务存在；（2）被告实际履行了部分付款义务，构成对协议的追认。',
            'unfavorable': '仅能证明已付款金额，不能证明欠款金额。需结合支付协议和结算单计算剩余欠款。',
            'opponent': '可能主张该付款是其他款项，与本案工程款无关；或主张已超额支付，要求返还。',
            'summary': '该证据是证明被告追认协议的关键，建议调取完整银行流水，固定付款事实。'
        },
        '结算单': {
            'beneficial': '工程款计算依据详细明确：（1）防水施工3679㎡×35元/㎡=128,765元；（2）点工47工日×400元/工日=18,800元；（3）扣减垃圾清理费2,000元；（4）合计145,565元。该结算单与支付协议金额一致，可互相印证。',
            'unfavorable': '（1）需核实结算单签字人身份和权限；（2）如被告否认结算单效力，需补充实际施工量证据；（3）点工部分缺乏详细工作内容记录。',
            'opponent': '（1）否认结算单签字人权限，主张结算无效；（2）质疑工程量和单价计算，要求重新核算；（3）主张存在质量问题应扣减工程款。',
            'summary': '该证据是计算工程款的重要依据，建议补充工程量确认单、验收记录等证据补强。'
        }
    }

    for key, value in analyses.items():
        if key in material_name:
            return value

    return {
        'beneficial': f'该材料与{case_data.get("案由", "本案")}相关，可作为证据支持我方主张。',
        'unfavorable': '需核实该证据的真实性、合法性、关联性。',
        'opponent': f'对方可能质疑该证据的真实性或关联性。',
        'summary': '该证据需与其他材料结合使用。'
    }

def generate_evidence_section(doc, case_data, evidence_num, evidence_name, evidence_type, screenshots):
    """生成单个证据的详细分析章节"""

    add_bold_label(doc, f'【证据{evidence_num}：{evidence_name}】')

    # 有利点分析
    add_bold_label(doc, '1. 有利点分析：')
    analysis = generate_detailed_analysis(case_data, evidence_type, 'pdf')
    add_normal_text(doc, analysis['beneficial'])
    doc.add_paragraph()

    # 关键内容截图
    add_bold_label(doc, '2. 关键内容截图：')
    if screenshots:
        for i, screenshot_path in enumerate(screenshots, 1):
            caption = f'图{evidence_num}-{i}：{evidence_name}关键内容'
            embed_image_in_doc(doc, screenshot_path, caption)
    else:
        add_normal_text(doc, '[未找到对应截图]')
    doc.add_paragraph()

    # 不利点
    add_bold_label(doc, '3. 不利点及风险：')
    add_normal_text(doc, analysis['unfavorable'])
    doc.add_paragraph()

    # 相对方攻击点
    add_bold_label(doc, '4. 相对方可能攻击点：')
    add_normal_text(doc, analysis['opponent'])
    doc.add_paragraph()

    # 分析小结
    add_bold_label(doc, '5. 分析小结：')
    add_normal_text(doc, analysis['summary'])
    doc.add_paragraph()
    doc.add_paragraph()

def generate_comprehensive_report(project_path, case_data, output_path=None):
    """生成完整详细的材料分析报告"""
    if output_path is None:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = Path(project_path) / f'材料分析报告_完整版_{timestamp}.docx'

    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # 扫描所有证据文件
    print("[*] 扫描证据文件...")
    evidence_files = scan_evidence_files(project_path)
    print(f"[*] 发现证据：")
    for etype, files in evidence_files.items():
        if files:
            print(f"    - {etype}: {len(files)}个文件")

    # 处理截图
    print("[*] 提取证据截图...")
    screenshot_dir = Path(project_path) / '截图证据'
    all_screenshots = process_evidence_screenshots(project_path, evidence_files, screenshot_dir)

    # ========== 标题 ==========
    add_title(doc, '案件材料分析报告', 22)

    plaintiff = case_data.get('原告公司名称', case_data.get('原告姓名', '原告'))
    defendant = case_data.get('被告公司名称', case_data.get('被告姓名', '被告'))
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(f'{plaintiff} vs {defendant}')
    set_chinese_font(subtitle_run, '宋体', 14)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.space_after = Pt(20)

    # ========== 一、案件基本信息 ==========
    add_heading1(doc, '一、案件基本信息')
    doc.add_paragraph()

    info_items = [
        ('案由', case_data.get('案由', '待补充')),
        ('委托人（原告）', plaintiff),
        ('对方当事人（被告）', defendant),
        ('受诉法院', case_data.get('受诉法院', '待补充')),
        ('案号', case_data.get('案号', '待补充')),
        ('承办律师', '待补充'),
        ('分析日期', datetime.now().strftime('%Y年%m月%d日')),
    ]

    for label, value in info_items:
        add_normal_text(doc, f'{label}：{value}')

    doc.add_paragraph()

    # ========== 二、工程项目背景 ==========
    add_heading1(doc, '二、工程项目背景')
    doc.add_paragraph()

    add_normal_text(doc, '根据现有材料，本案涉及的工程项目信息如下：', bold=True)
    doc.add_paragraph()

    background = (
        f'合同关系链：{plaintiff}（防水工程分包方）→ {defendant}（净化工程总包方之一）→ '
        f'龙港市人民医院项目（业主）。根据中标结果公告，项目中标价约7855万元，工期180天。\n\n'
        f'工程内容：龙港市人民医院净化工程防水施工，包括防水施工3679㎡、点工47工日等。\n\n'
        f'合同履行情况：双方于2025年12月22日签订《支付协议》，确认工程款145,565元，'
        f'约定分期支付。被告已支付50,000元，剩余95,565元未付。'
    )
    add_normal_text(doc, background)
    doc.add_paragraph()

    # ========== 三、证据材料清单与分析 ==========
    add_heading1(doc, '三、证据材料清单与分析')

    # （一）材料清单
    add_heading2(doc, '（一）材料清单')
    doc.add_paragraph()

    total_files = sum(len(files) for files in evidence_files.values())
    add_normal_text(doc, f'本次共扫描到{total_files}份证据材料：')
    doc.add_paragraph()

    evidence_num = 1
    for etype, files in evidence_files.items():
        if files:
            for f in files:
                add_normal_text(doc, f'{evidence_num}. {f.name}（{f.suffix[1:].upper()}，{f.stat().st_size / 1024:.1f} KB）')
                evidence_num += 1

    doc.add_paragraph()

    # （二）逐份证据详细分析
    add_heading2(doc, '（二）逐份证据详细分析')
    doc.add_paragraph()

    # 证据1：中标公告
    if evidence_files['中标公告']:
        generate_evidence_section(doc, case_data, 1, '中标结果公告', '中标公告', all_screenshots.get('中标公告', []))

    # 证据2：支付协议
    if evidence_files['支付协议']:
        generate_evidence_section(doc, case_data, 2, '支付协议（2025年12月22日）', '支付协议', all_screenshots.get('支付协议', []))

    # 证据3：银行回单
    if evidence_files['银行回单']:
        generate_evidence_section(doc, case_data, 3, '银行电子回单×2', '银行回单', all_screenshots.get('银行回单', []))

    # 证据4：结算单
    if evidence_files['结算单']:
        generate_evidence_section(doc, case_data, 4, '劳务分包结算单（2024年12月27日）', '结算单', all_screenshots.get('结算单', []))

    # （三）欠款金额计算
    add_heading2(doc, '（三）欠款金额计算')
    doc.add_paragraph()

    calculations = [
        '1. 总工程款：145,565元（根据支付协议和结算单确认）',
        '2. 已付款：50,000元（银行电子回单×2证明）',
        '   - 2026年1月30日付款：20,000元',
        '   - 2026年2月14日付款：30,000元',
        '3. 剩余未付款：95,565元（计算：145,565 - 50,000 = 95,565）',
        '4. 逾期利息：以95,565元为基数，自2026年2月17日起按LPR计算至实际付清之日止',
    ]

    for calc in calculations:
        add_normal_text(doc, calc)
    doc.add_paragraph()

    # ========== 四、核心争议焦点 ==========
    add_heading1(doc, '四、核心争议焦点')
    doc.add_paragraph()

    add_risk_section(doc, '高风险', '项目经理签字效力问题',
        f'支付协议由乙方经办人"翟振"签字，但未见授权委托书或职务证明。{defendant}后续付款行为可视为对协议的追认，但仍需补强翟振的身份及权限证明。')

    add_risk_section(doc, '中风险', '协议主体瑕疵问题',
        f'签约乙方为"龙港市人民医院净化工程项目部"，该主体并非独立法人。虽然{defendant}实际履行了部分付款义务，但对方仍可能以此抗辩协议效力。')

    add_risk_section(doc, '中风险', '剩余款项支付条件争议',
        '协议第二条约定"剩余部分甲乙双方协商处理"，该表述模糊，对方可能主张剩余款项支付条件尚未成就。')

    add_bold_label(doc, '4. 已付款金额认定')
    add_normal_text(doc, '根据银行回单，千业公司已实际支付50,000元，与协议约定的前两期付款总额95,565元不符，存在45,565元缺口。')
    doc.add_paragraph()

    # ========== 五、风险提示 ==========
    add_heading1(doc, '五、风险提示')

    add_risk_section(doc, '高风险', '项目经理签字效力',
        '如无法证明翟振具有代表千业公司签署付款协议的权限，对方可能否认协议效力。建议立即收集：翟振的授权委托书、职务证明、社保缴纳记录等。')

    add_risk_section(doc, '中风险', '协议主体瑕疵',
        f'签约主体为项目部，存在主体不适格风险。{defendant}后续付款行为可作为追认证据。')

    add_risk_section(doc, '中风险', '剩余款项抗辩',
        '协议"剩余部分协商处理"的表述可能被对方利用。建议收集协商过程的证据（微信记录、通话录音等）。')

    add_risk_section(doc, '低风险', '诉讼时效',
        '第二期付款期限为2026年2月16日，目前仍在诉讼时效内。但建议尽快采取法律行动。')

    # ========== 六、证据缺口与补强建议 ==========
    add_heading1(doc, '六、证据缺口与补强建议')
    doc.add_paragraph()

    gaps = [
        ('建设工程施工合同', '需收集原件，证明合同关系及工程内容'),
        ('项目经理身份证明及授权文件', '关键：证明"翟振"签字效力'),
        ('工程验收/结算文件', '除结算单外，需收集竣工验收文件等'),
        ('催款记录', '证明时效及对方违约'),
        ('损失计算依据', '逾期利息、违约金等计算依据'),
    ]

    for i, (gap_name, gap_desc) in enumerate(gaps, 1):
        add_bold_label(doc, f'{i}. {gap_name}')
        add_normal_text(doc, gap_desc)
        doc.add_paragraph()

    # ========== 七、证据链分析 ==========
    add_heading1(doc, '七、证据链分析')
    doc.add_paragraph()

    chain_analysis = (
        f'现有{total_files}份证据材料已形成初步证据链：\n\n'
        f'1. 【中标结果公告】→ 证明项目背景和被告主体资格\n'
        f'2. 【支付协议】→ 确认债务金额和付款安排（核心证据）\n'
        f'3. 【银行电子回单】→ 证明被告部分履行，构成追认\n'
        f'4. 【劳务分包结算单】→ 印证工程款计算依据\n\n'
        f'证据链完整性评估：现有证据已能证明债务存在和欠款金额，但签字效力和主体瑕疵问题需要补强。'
    )
    add_normal_text(doc, chain_analysis)
    doc.add_paragraph()

    # ========== 八、综合结论 ==========
    add_heading1(doc, '八、综合结论')
    doc.add_paragraph()

    conclusion = (
        f'经对{plaintiff}与{defendant}建设工程合同纠纷案件现有材料进行详细分析，共梳理出{total_files}份关键证据材料。'
        f'已确认总工程款145,565元，已付款50,000元，剩余95,565元未付。\n\n'
        f'【优势分析】\n'
        f'1. 支付协议明确了债务金额，是最直接的债权凭证\n'
        f'2. 被告已部分履行（付款50,000元），构成对协议的追认\n'
        f'3. 银行回单和结算单可互相印证，证据链较为完整\n\n'
        f'【风险提示】\n'
        f'1. 项目经理签字效力和协议主体瑕疵是主要风险点\n'
        f'2. "剩余部分协商处理"的表述可能被对方利用抗辩\n\n'
        f'【建议方案】\n'
        f'建议尽快补强关键证据后，及时向龙港市人民法院提起诉讼。'
    )

    add_normal_text(doc, conclusion)

    # 保存
    doc.save(str(output_path))
    print(f"[✓] 材料分析报告（完整版）已生成: {output_path}")
    return str(output_path)

if __name__ == '__main__':
    import argparse

    parser = argparse.ArgumentParser(description='材料分析工具（v3.0 智能截图版）')
    parser.add_argument('--project', '-p', required=True, help='项目路径')
    parser.add_argument('--data', '-d', required=True, help='案件数据JSON文件路径')
    parser.add_argument('--output', '-o', help='输出报告路径（可选）')

    args = parser.parse_args()

    with open(args.data, 'r', encoding='utf-8') as f:
        case_data = json.load(f)

    report_path = generate_comprehensive_report(args.project, case_data, args.output)

    print(f"\n[*] 完整分析报告生成完成！")
    print(f"[*] 报告路径: {report_path}")
