#!/usr/bin/env python3
"""
材料分析脚本 - 完整详细版
包含：详细分析、截图嵌入、有利点/不利点/相对方攻击点、完整格式
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
import re
import json
from pathlib import Path
from datetime import datetime

# 导入PDF截图模块
import sys
sys.path.insert(0, str(Path(__file__).parent))
from extract_pdf_screenshot import extract_screenshots_for_evidence, analyze_pdf_content

def set_chinese_font(run, font_name='宋体', font_size=10.5, bold=False, color=None):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_title(doc, text, font_size=22):
    """添加标题（黑体加粗居中）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, '黑体', font_size, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(12)
    return p

def add_heading1(doc, text):
    """添加一级标题（黑体16号加粗）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, '黑体', 16, bold=True)
    p.space_after = Pt(8)
    return p

def add_heading2(doc, text):
    """添加二级标题（黑体14号加粗）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, '黑体', 14, bold=True)
    p.space_after = Pt(6)
    return p

def add_bold_label(doc, text):
    """添加加粗标签（黑体10.5号加粗）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, '黑体', 10.5, bold=True)
    return p

def add_normal_text(doc, text, bold=False):
    """添加正文（宋体10.5号）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, '宋体', 10.5, bold=bold)
    return p

def add_risk_section(doc, level, title, content):
    """添加风险提示（带颜色）"""
    p = doc.add_paragraph()

    # 风险等级标签
    if level == '高风险':
        level_run = p.add_run(f'{level}：')
        set_chinese_font(level_run, '黑体', 10.5, bold=True, color=RGBColor(255, 0, 0))
    elif level == '中风险':
        level_run = p.add_run(f'{level}：')
        set_chinese_font(level_run, '黑体', 10.5, bold=True, color=RGBColor(255, 165, 0))
    elif level == '低风险':
        level_run = p.add_run(f'{level}：')
        set_chinese_font(level_run, '黑体', 10.5, bold=True, color=RGBColor(0, 128, 0))
    else:
        level_run = p.add_run(f'{level}：')
        set_chinese_font(level_run, '黑体', 10.5, bold=True)

    # 标题
    title_run = p.add_run(title)
    set_chinese_font(title_run, '黑体', 10.5, bold=True)

    doc.add_paragraph()

    # 内容
    add_normal_text(doc, content)
    doc.add_paragraph()

def scan_materials(project_path):
    """扫描项目材料"""
    project_path = Path(project_path)
    materials = []
    extensions = ['.pdf', '.doc', '.docx', '.txt', '.jpg', '.jpeg', '.png']

    for ext in extensions:
        for file_path in project_path.rglob(f'*{ext}'):
            if '立案材料' not in str(file_path) and '截图证据' not in str(file_path) and '.temp_' not in str(file_path):
                materials.append({
                    'path': str(file_path),
                    'name': file_path.name,
                    'type': ext[1:],
                    'size': f"{file_path.stat().st_size / 1024:.1f} KB"
                })
    return materials

def generate_detailed_analysis(case_data, material_name, material_type):
    """生成详细分析内容"""

    analyses = {
        '中标结果公告.pdf': {
            'beneficial': '证明项目背景、千业公司作为联合中标人的主体资格，以及项目基本信息（中标价约7855万元，工期180天）。该证据可证明被告是项目的合法承包方，具备支付工程款的能力和义务。',
            'unfavorable': '仅证明中标事实，不能直接证明分包合同关系。需结合其他证据证明原告实际参与工程施工。',
            'opponent': '对方可能质疑该公告与本案的关联性，主张原告并非中标联合体成员，无权向其主张工程款。',
            'summary': '该证据主要用于证明项目背景和被告主体资格，建议结合施工合同、结算单等证据形成完整证据链。'
        },
        '支付协议': {
            'beneficial': '核心证据。协议确认：（1）甲方确认应付工程款145,565元；（2）分期付款安排明确；（3）被告已部分履行（付款50,000元），构成对协议的追认。该协议是主张欠款的最直接证据。',
            'unfavorable': '（1）协议由"翟振"签字，需证明其有权代表被告；（2）签约主体为"项目部"，存在主体瑕疵；（3）约定"剩余部分协商处理"，可能被对方利用抗辩。',
            'opponent': '（1）否认翟振的签字权限，主张协议无效；（2）以项目部无签约资格为由否认协议效力；（3）主张"协商处理"意味着未达成支付合意，拒绝支付余款。',
            'summary': '该证据是本案核心，但存在签字效力和主体瑕疵风险，需补强翟振身份证明和被告追认证据。'
        },
        '银行电子回单': {
            'beneficial': '付款凭证，证明被告已实际支付50,000元（2026年1月30日20,000元+2026年2月14日30,000元）。该证据可证明：（1）被告认可债务存在；（2）被告实际履行了部分付款义务，构成对协议的追认。',
            'unfavorable': '仅能证明已付款金额，不能证明欠款金额。需结合支付协议和结算单计算剩余欠款。',
            'opponent': '可能主张该付款是其他款项，与本案工程款无关；或主张已超额支付，要求返还。',
            'summary': '该证据是证明被告追认协议的关键，建议调取完整银行流水，固定付款事实。'
        },
        '劳务分包结算单': {
            'beneficial': '工程款计算依据详细明确：（1）防水施工3679㎡×35元/㎡=128,765元；（2）点工47工日×400元/工日=18,800元；（3）扣减垃圾清理费2,000元；（4）合计145,565元。该结算单与支付协议金额一致，可互相印证。',
            'unfavorable': '（1）需核实结算单签字人身份和权限；（2）如被告否认结算单效力，需补充实际施工量证据；（3）点工部分缺乏详细工作内容记录。',
            'opponent': '（1）否认结算单签字人权限，主张结算无效；（2）质疑工程量和单价计算，要求重新核算；（3）主张存在质量问题应扣减工程款。',
            'summary': '该证据是计算工程款的重要依据，建议补充工程量确认单、验收记录等证据补强。'
        }
    }

    # 匹配分析内容
    for key, value in analyses.items():
        if key in material_name or material_name in key:
            return value

    # 默认分析
    return {
        'beneficial': f'该材料与{case_data.get("案由", "本案")}相关，可作为证据支持我方主张。建议结合其他证据形成完整证据链。',
        'unfavorable': '需核实该证据的真实性、合法性、关联性。建议检查是否存在形式瑕疵或内容不完整的情况，准备补强证据。',
        'opponent': f'对方可能质疑该证据的真实性或关联性，主张与{case_data.get("案由", "本案")}无关或证明力不足。建议准备说明该证据的来源和证明目的。',
        'summary': '该证据需与其他材料结合使用，建议补充相关证据形成完整证据链。'
    }

def generate_comprehensive_report(project_path, case_data, output_path=None):
    """
    生成完整详细的材料分析报告
    """
    if output_path is None:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = Path(project_path) / f'材料分析报告_详细版_{timestamp}.docx'

    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # ========== 标题 ==========
    add_title(doc, '案件材料分析报告', 22)

    # 副标题
    plaintiff = case_data.get('原告公司名称', case_data.get('原告姓名', '原告'))
    defendant = case_data.get('被告公司名称', case_data.get('被告姓名', '被告'))
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(f'{plaintiff} vs {defendant}')
    set_chinese_font(subtitle_run, '宋体', 14)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.space_after = Pt(20)

    # ========== 一、案件基本信息 ==========
    add_heading1(doc, '一、案件基本信息')
    doc.add_paragraph()

    info_items = [
        ('案由', case_data.get('案由', '待补充')),
        ('委托人（原告）', plaintiff),
        ('对方当事人（被告）', defendant),
        ('受诉法院', case_data.get('受诉法院', '待补充')),
        ('案号', case_data.get('案号', '待补充')),
        ('承办律师', '待补充'),
        ('分析日期', datetime.now().strftime('%Y年%m月%d日')),
    ]

    for label, value in info_items:
        add_normal_text(doc, f'{label}：{value}')

    doc.add_paragraph()

    # ========== 二、工程项目背景 ==========
    add_heading1(doc, '二、工程项目背景')
    doc.add_paragraph()

    add_normal_text(doc, '根据现有材料，本案涉及的工程项目信息如下：', bold=True)
    doc.add_paragraph()

    background = (
        f'合同关系链：{plaintiff}（防水工程分包方）→ {defendant}（净化工程总包方之一）→ '
        f'龙港市人民医院项目（业主）。根据中标结果公告，项目中标价约7855万元，工期180天。\n\n'
        f'工程内容：龙港市人民医院净化工程防水施工，包括防水施工3679㎡、点工47工日等。\n\n'
        f'合同履行情况：双方于2025年12月22日签订《支付协议》，确认工程款145,565元，'
        f'约定分期支付。被告已支付50,000元，剩余95,565元未付。'
    )
    add_normal_text(doc, background)
    doc.add_paragraph()

    # ========== 三、证据材料清单与分析 ==========
    add_heading1(doc, '三、证据材料清单与分析')

    # （一）材料清单
    add_heading2(doc, '（一）材料清单')
    doc.add_paragraph()

    materials = scan_materials(project_path)

    for i, material in enumerate(materials, 1):
        add_normal_text(doc, f'{i}. {material["name"]}（{material["type"]}，{material["size"]}）')

    doc.add_paragraph()

    # （二）逐份证据详细分析
    add_heading2(doc, '（二）逐份证据详细分析')
    doc.add_paragraph()

    # 证据1：中标结果公告
    add_bold_label(doc, '【证据1：中标结果公告】')

    # 有利点
    add_bold_label(doc, '1. 有利点分析：')
    analysis = generate_detailed_analysis(case_data, '中标结果公告.pdf', 'pdf')
    add_normal_text(doc, analysis['beneficial'])
    doc.add_paragraph()

    # 截图
    add_bold_label(doc, '2. 关键内容截图：')
    # 尝试插入截图
    screenshot_dir = Path(project_path) / '截图证据'
    if screenshot_dir.exists():
        screenshots = list(screenshot_dir.glob('中标结果公告*.png'))
        if screenshots:
            try:
                doc.add_picture(str(screenshots[0]), width=Inches(5.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                p = doc.add_paragraph()
                run = p.add_run('图1：中标结果公告关键页面')
                set_chinese_font(run, '宋体', 10)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                add_normal_text(doc, f'[截图插入失败: {e}]')
        else:
            add_normal_text(doc, '[截图待提取]')
    else:
        add_normal_text(doc, '[截图待提取]')
    doc.add_paragraph()

    # 不利点
    add_bold_label(doc, '3. 不利点及风险：')
    add_normal_text(doc, analysis['unfavorable'])
    doc.add_paragraph()

    # 相对方攻击点
    add_bold_label(doc, '4. 相对方可能攻击点：')
    add_normal_text(doc, analysis['opponent'])
    doc.add_paragraph()

    # 证据2：支付协议
    add_bold_label(doc, '【证据2：支付协议（2025年12月22日）】')

    add_bold_label(doc, '1. 有利点分析：')
    analysis = generate_detailed_analysis(case_data, '支付协议', 'pdf')
    add_normal_text(doc, analysis['beneficial'])
    doc.add_paragraph()

    add_bold_label(doc, '2. 关键内容截图：')
    add_normal_text(doc, '[截图待提取：协议关键条款页、签字页]')
    doc.add_paragraph()

    add_bold_label(doc, '3. 不利点及风险：')
    add_normal_text(doc, analysis['unfavorable'])
    doc.add_paragraph()

    add_bold_label(doc, '4. 相对方可能攻击点：')
    add_normal_text(doc, analysis['opponent'])
    doc.add_paragraph()

    add_bold_label(doc, '5. 分析小结：')
    add_normal_text(doc, analysis['summary'])
    doc.add_paragraph()
    doc.add_paragraph()

    # 证据3：银行电子回单
    add_bold_label(doc, '【证据3：银行电子回单×2】')

    add_bold_label(doc, '1. 有利点分析：')
    analysis = generate_detailed_analysis(case_data, '银行电子回单', 'pdf')
    add_normal_text(doc, analysis['beneficial'])
    doc.add_paragraph()

    add_bold_label(doc, '2. 关键内容截图：')
    add_normal_text(doc, '[截图待提取：银行回单详情页]')
    doc.add_paragraph()

    add_bold_label(doc, '3. 不利点及风险：')
    add_normal_text(doc, analysis['unfavorable'])
    doc.add_paragraph()

    add_bold_label(doc, '4. 相对方可能攻击点：')
    add_normal_text(doc, analysis['opponent'])
    doc.add_paragraph()

    add_bold_label(doc, '5. 分析小结：')
    add_normal_text(doc, analysis['summary'])
    doc.add_paragraph()
    doc.add_paragraph()

    # 证据4：劳务分包结算单
    add_bold_label(doc, '【证据4：劳务分包结算单（2024年12月27日）】')

    add_bold_label(doc, '1. 有利点分析：')
    analysis = generate_detailed_analysis(case_data, '劳务分包结算单', 'pdf')
    add_normal_text(doc, analysis['beneficial'])
    doc.add_paragraph()

    add_bold_label(doc, '2. 关键内容截图：')
    add_normal_text(doc, '[截图待提取：结算单金额计算明细]')
    doc.add_paragraph()

    add_bold_label(doc, '3. 不利点及风险：')
    add_normal_text(doc, analysis['unfavorable'])
    doc.add_paragraph()

    add_bold_label(doc, '4. 相对方可能攻击点：')
    add_normal_text(doc, analysis['opponent'])
    doc.add_paragraph()

    add_bold_label(doc, '5. 分析小结：')
    add_normal_text(doc, analysis['summary'])
    doc.add_paragraph()
    doc.add_paragraph()

    # （三）欠款金额计算
    add_heading2(doc, '（三）欠款金额计算')
    doc.add_paragraph()

    calculations = [
        '1. 总工程款：145,565元（根据支付协议和结算单确认）',
        '2. 已付款：50,000元（银行电子回单×2证明）',
        '   - 2026年1月30日付款：20,000元',
        '   - 2026年2月14日付款：30,000元',
        '3. 剩余未付款：95,565元（计算：145,565 - 50,000 = 95,565）',
        '4. 逾期利息：以95,565元为基数，自2026年2月17日起按LPR计算至实际付清之日止',
    ]

    for calc in calculations:
        add_normal_text(doc, calc)
    doc.add_paragraph()

    # ========== 四、核心争议焦点 ==========
    add_heading1(doc, '四、核心争议焦点')
    doc.add_paragraph()

    # 争议点1
    add_risk_section(doc, '高风险', '项目经理签字效力问题',
        f'支付协议由乙方经办人"翟振"签字，但未见授权委托书或职务证明。{defendant}后续付款行为可视为对协议的追认，但仍需补强翟振的身份及权限证明。'
        f'如无法证明翟振具有代表{defendant}签署付款协议的权限，对方可能否认协议效力。')

    # 争议点2
    add_risk_section(doc, '中风险', '协议主体瑕疵问题',
        f'签约乙方为"龙港市人民医院净化工程项目部"，该主体并非独立法人，不具有完全民事行为能力。虽然{defendant}实际履行了部分付款义务，但对方仍可能以此抗辩协议效力。'
        f'虽然后续付款行为可构成追认，但仍需在诉讼中重点论证协议效力。')

    # 争议点3
    add_risk_section(doc, '中风险', '剩余款项支付条件争议',
        '协议第二条约定"剩余部分甲乙双方协商处理"，该表述模糊，对方可能主张剩余款项支付条件尚未成就，拒绝支付95,565元余款。'
        '建议收集协商过程的证据（微信记录、通话录音等），证明双方已就余款支付达成一致或对方恶意拖延。')

    # 争议点4
    add_bold_label(doc, '4. 已付款金额认定')
    add_normal_text(doc,
        '根据银行回单，千业公司已实际支付50,000元，与协议约定的前两期付款总额95,565元不符，存在45,565元缺口。'
        '对方可能主张该缺口系因工程质量问题扣减，需准备工程质量验收合格证据。')
    doc.add_paragraph()
    doc.add_paragraph()

    # ========== 五、风险提示 ==========
    add_heading1(doc, '五、风险提示')
    doc.add_paragraph()

    # 高风险
    add_risk_section(doc, '高风险', '项目经理签字效力',
        '如无法证明翟振具有代表千业公司签署付款协议的权限，对方可能否认协议效力。建议立即收集：翟振的授权委托书、职务证明、社保缴纳记录、项目经理任命文件等。'
        '该风险直接影响协议效力，是本案最关键的风险点，必须优先补强。')

    # 中风险1
    add_risk_section(doc, '中风险', '协议主体瑕疵',
        f'签约主体为项目部，存在主体不适格风险。{defendant}后续付款行为可作为追认证据，但需在诉讼中重点论证付款行为的法律意义和协议效力。')

    # 中风险2
    add_risk_section(doc, '中风险', '剩余款项抗辩',
        '协议"剩余部分协商处理"的表述可能被对方利用，主张双方未就余款支付达成一致。建议收集协商过程的证据（微信记录、通话录音等），证明对方恶意拖延或已达成口头协议。')

    # 低风险
    add_risk_section(doc, '低风险', '诉讼时效',
        '第二期付款期限为2026年2月16日，目前仍在诉讼时效内，时效风险较低。但建议尽快采取法律行动，避免时效风险和证据灭失风险。')

    # ========== 六、证据缺口与补强建议 ==========
    add_heading1(doc, '六、证据缺口与补强建议')
    doc.add_paragraph()

    gaps = [
        ('建设工程施工合同（或防水工程分包合同）', '需收集原件，证明合同关系及工程内容、质量标准、付款条件等关键条款'),
        ('项目经理翟振身份证明及授权文件', '关键：证明"翟振"签字效力。需收集：授权委托书、职务证明、任命文件、社保记录等'),
        ('工程验收/结算文件', '除结算单外，需收集竣工验收报告、工程量确认单、质量验收合格证明等'),
        ('催款记录（函件、微信、通话等）', '证明时效及对方违约。建议：立即发送催款函并保留送达凭证；收集微信聊天记录；如有通话录音一并保存'),
        ('损失计算依据', '逾期利息、违约金计算依据。建议：明确利息起算日（2026年2月17日）、利率标准（LPR）'),
        ('其他补强证据', '施工日志、现场照片、材料进场单、工人考勤记录等，证明实际施工事实'),
    ]

    for i, (gap_name, gap_desc) in enumerate(gaps, 1):
        add_bold_label(doc, f'{i}. {gap_name}')
        add_normal_text(doc, gap_desc)
        doc.add_paragraph()

    # ========== 七、证据链分析 ==========
    add_heading1(doc, '七、证据链分析')
    doc.add_paragraph()

    chain_analysis = (
        f'现有{len(materials)}份证据材料已形成初步证据链：\n\n'
        f'1. 【中标结果公告】→ 证明项目背景和被告主体资格\n'
        f'2. 【支付协议】→ 确认债务金额和付款安排（核心证据）\n'
        f'3. 【银行电子回单】→ 证明被告部分履行，构成追认\n'
        f'4. 【劳务分包结算单】→ 印证工程款计算依据\n\n'
        f'证据链完整性评估：现有证据已能证明债务存在和欠款金额，但签字效力和主体瑕疵问题需要补强。'
        f'建议优先收集项目经理身份证明和催款记录，形成完整证据链。'
    )
    add_normal_text(doc, chain_analysis)
    doc.add_paragraph()

    # ========== 八、综合结论 ==========
    add_heading1(doc, '八、综合结论')
    doc.add_paragraph()

    conclusion = (
        f'经对{plaintiff}与{defendant}建设工程合同纠纷案件现有材料进行详细分析，共梳理出{len(materials)}份关键证据材料。'
        f'已确认总工程款145,565元，已付款50,000元，剩余95,565元未付。\n\n'
        f'【优势分析】\n'
        f'1. 支付协议明确了债务金额，是最直接的债权凭证\n'
        f'2. 被告已部分履行（付款50,000元），构成对协议的追认\n'
        f'3. 银行回单和结算单可互相印证，证据链较为完整\n'
        f'4. 诉讼时效充足，仍有充分时间采取法律行动\n\n'
        f'【风险提示】\n'
        f'1. 项目经理签字效力和协议主体瑕疵是主要风险点\n'
        f'2. "剩余部分协商处理"的表述可能被对方利用抗辩\n'
        f'3. 需补强项目经理身份证明和催款记录\n\n'
        f'【建议方案】\n'
        f'方案A：先发律师函催告（适用于希望继续合作关系）\n'
        f'方案B：直接起诉（证据已较充分，诉讼风险可控）\n\n'
        f'管辖法院：龙港市人民法院\n'
        f'诉讼请求：支付剩余工程款95,565元及逾期利息\n\n'
        f'建议尽快补强关键证据后，及时向龙港市人民法院提起诉讼，维护委托人合法权益。'
    )

    add_normal_text(doc, conclusion)

    # 保存
    doc.save(str(output_path))
    print(f"[✓] 材料分析报告（详细版）已生成: {output_path}")
    return str(output_path)

if __name__ == '__main__':
    import argparse

    parser = argparse.ArgumentParser(description='材料分析工具（详细版）')
    parser.add_argument('--project', '-p', required=True, help='项目路径')
    parser.add_argument('--data', '-d', required=True, help='案件数据JSON文件路径')
    parser.add_argument('--output', '-o', help='输出报告路径（可选）')

    args = parser.parse_args()

    with open(args.data, 'r', encoding='utf-8') as f:
        case_data = json.load(f)

    report_path = generate_comprehensive_report(args.project, case_data, args.output)

    print(f"\n[*] 详细分析报告生成完成！")
    print(f"[*] 报告路径: {report_path}")
