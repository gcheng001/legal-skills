#!/usr/bin/env python3
"""
基于现有模板创建填空版模板
读取现有docx文件并添加{{字段名}}占位符
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
import os

def set_chinese_font(run, font_name='仿宋', font_size=14):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def extract_text_from_docx(file_path):
    """提取docx文件中的所有文本"""
    try:
        doc = Document(file_path)
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text.strip())
        return '\n'.join(text), doc
    except Exception as e:
        return f"Error: {e}", None

def create_enhanced_complaint_template():
    """
    基于标准民事起诉状格式创建填空版模板
    包含完整的起诉状结构和所有必要字段
    """
    doc = Document()

    # 设置中文字体支持
    style = doc.styles['Normal']
    style.font.name = '仿宋'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    style.font.size = Pt(14)

    # 标题
    title = doc.add_paragraph()
    title_run = title.add_run('民事起诉状')
    set_chinese_font(title_run, '黑体', 22)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(12)

    # 原告信息
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('原告：{{原告姓名}}')
    set_chinese_font(run, '黑体', 14)

    p = doc.add_paragraph()
    run = p.add_run('性别：{{原告性别}}　　民族：{{原告民族}}　　出生日期：{{原告出生日期}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('身份证号：{{原告身份证号}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('住所地：{{原告地址}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('联系电话：{{原告电话}}')
    set_chinese_font(run, '仿宋', 14)

    # 如果原告是公司
    p = doc.add_paragraph()
    run = p.add_run('（或）原告：{{原告公司名称}}')
    set_chinese_font(run, '黑体', 14)

    p = doc.add_paragraph()
    run = p.add_run('统一社会信用代码：{{原告公司代码}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('住所地：{{原告公司地址}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('法定代表人：{{原告法定代表人}}　　职务：{{原告法定代表人职务}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('联系电话：{{原告公司电话}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    # 被告信息
    p = doc.add_paragraph()
    run = p.add_run('被告：{{被告姓名}}')
    set_chinese_font(run, '黑体', 14)

    p = doc.add_paragraph()
    run = p.add_run('性别：{{被告性别}}　　民族：{{被告民族}}　　出生日期：{{被告出生日期}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('身份证号：{{被告身份证号}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('住所地：{{被告地址}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('联系电话：{{被告电话}}')
    set_chinese_font(run, '仿宋', 14)

    # 如果被告是公司
    p = doc.add_paragraph()
    run = p.add_run('（或）被告：{{被告公司名称}}')
    set_chinese_font(run, '黑体', 14)

    p = doc.add_paragraph()
    run = p.add_run('统一社会信用代码：{{被告公司代码}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('住所地：{{被告公司地址}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('法定代表人：{{被告法定代表人}}　　职务：{{被告法定代表人职务}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('联系电话：{{被告公司电话}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    # 案由
    p = doc.add_paragraph()
    run = p.add_run('案由：{{案由}}')
    set_chinese_font(run, '黑体', 14)

    doc.add_paragraph()

    # 诉讼请求
    p = doc.add_paragraph()
    run = p.add_run('诉讼请求：')
    set_chinese_font(run, '黑体', 14)

    p = doc.add_paragraph()
    run = p.add_run('1、{{诉讼请求1}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('2、{{诉讼请求2}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('3、{{诉讼请求3}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('4、{{诉讼请求4}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('5、{{诉讼请求5}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    # 事实与理由
    p = doc.add_paragraph()
    run = p.add_run('事实与理由：')
    set_chinese_font(run, '黑体', 14)

    p = doc.add_paragraph()
    run = p.add_run('{{事实与理由}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    # 证据清单
    p = doc.add_paragraph()
    run = p.add_run('证据清单：')
    set_chinese_font(run, '黑体', 14)

    p = doc.add_paragraph()
    run = p.add_run('{{证据清单}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    # 此致
    p = doc.add_paragraph()
    run = p.add_run('此致')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('{{受诉法院}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()
    doc.add_paragraph()

    # 具状人和日期（右对齐）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('具状人(签字、盖章)：')
    set_chinese_font(run, '仿宋', 14)
    run = p.add_run('{{具状人}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('日期：')
    set_chinese_font(run, '仿宋', 14)
    run = p.add_run('{{日期}}')
    set_chinese_font(run, '仿宋', 14)

    return doc

def create_enhanced_poa_template():
    """
    创建完整的委托手续模板
    包含：授权委托书、律师事务所函、律师证复印件说明
    """
    doc = Document()

    # 设置中文字体
    style = doc.styles['Normal']
    style.font.name = '仿宋'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    style.font.size = Pt(14)

    # ========== 第一部分：授权委托书 ==========
    title = doc.add_paragraph()
    title_run = title.add_run('授 权 委 托 书')
    set_chinese_font(title_run, '黑体', 22)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(12)

    doc.add_paragraph()

    # 委托人信息
    p = doc.add_paragraph()
    run = p.add_run('委托人：{{原告姓名}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('身份证号/统一社会信用代码：{{原告身份证号}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('住址/住所地：{{原告地址}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('联系电话：{{原告电话}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    # 受托人信息
    p = doc.add_paragraph()
    run = p.add_run('受托人：{{代理人姓名}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('工作单位：{{律所名称}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('职务：律师')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('联系电话：{{代理人电话}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('执业证号：{{代理人执业证号}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    # 委托事项
    p = doc.add_paragraph()
    run = p.add_run('现委托{{代理人姓名}}在{{原告姓名}}与{{被告姓名}}{{案由}}纠纷一案中，作为我方委托代理人。')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('委托事项和权限如下：')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('{{委托权限}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('（如：特别授权：代为起诉、应诉、承认、放弃、变更诉讼请求，进行和解、调解，代收法律文书等）')
    set_chinese_font(run, '仿宋', 12)

    doc.add_paragraph()
    doc.add_paragraph()

    # 签名区
    p = doc.add_paragraph()
    run = p.add_run('委托人(签字/盖章)：')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('日期：{{日期}}')
    set_chinese_font(run, '仿宋', 14)

    # 分页
    doc.add_page_break()

    # ========== 第二部分：律师事务所函 ==========
    title = doc.add_paragraph()
    title_run = title.add_run('律 师 事 务 所 函')
    set_chinese_font(title_run, '黑体', 22)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('{{受诉法院}}：')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('本所接受{{原告姓名}}的委托，指派{{代理人姓名}}律师担任其与你院受理的（{{案由}}）纠纷一案中{{原告姓名}}的委托代理人。')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('特此函告')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('附：1、授权委托书一份')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('　　2、律师证复印件一份')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()
    doc.add_paragraph()

    # 律所盖章区（右对齐）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('{{律所名称}}（盖章）')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('日期：{{日期}}')
    set_chinese_font(run, '仿宋', 14)

    return doc

def create_enhanced_legal_representative_template():
    """
    创建法定代表人身份证明模板
    """
    doc = Document()

    # 设置中文字体
    style = doc.styles['Normal']
    style.font.name = '仿宋'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    style.font.size = Pt(14)

    # 标题
    title = doc.add_paragraph()
    title_run = title.add_run('法定代表人身份证明书')
    set_chinese_font(title_run, '黑体', 22)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(12)

    doc.add_paragraph()

    # 正文
    p = doc.add_paragraph()
    run = p.add_run('{{原告公司名称}}（单位名称）在我单位任{{法定代表人姓名}}职务，系我单位的法定代表人，特此证明。')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('附：法定代表人身份证复印件')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()
    doc.add_paragraph()

    # 单位盖章区（右对齐）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('单位（盖章）：{{原告公司名称}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('日期：{{日期}}')
    set_chinese_font(run, '仿宋', 14)

    return doc

if __name__ == '__main__':
    base_path = '~/templates/civil'

    # 创建增强版起诉状模板
    complaint_doc = create_enhanced_complaint_template()
    complaint_path = os.path.join(base_path, '起诉状', '民事起诉状-通用模板.docx')
    complaint_doc.save(complaint_path)
    print(f'✓ 已更新: {complaint_path}')

    # 创建增强版委托手续模板
    poa_doc = create_enhanced_poa_template()
    poa_path = os.path.join(base_path, '委托手续-填空模板.docx')
    poa_doc.save(poa_path)
    print(f'✓ 已更新: {poa_path}')

    # 创建增强版法定代表人身份证明模板
    lr_doc = create_enhanced_legal_representative_template()
    lr_path = os.path.join(base_path, '法定代表人身份证明-填空模板.docx')
    lr_doc.save(lr_path)
    print(f'✓ 已更新: {lr_path}')

    print('\n三个填空版模板已更新完成！')
    print('\n支持的占位符包括：')
    print('原告相关: {{原告姓名}}, {{原告性别}}, {{原告身份证号}}, {{原告地址}}, {{原告电话}}, {{原告公司名称}}, {{原告公司代码}}, {{原告法定代表人}}, {{原告法定代表人职务}}')
    print('被告相关: {{被告姓名}}, {{被告性别}}, {{被告身份证号}}, {{被告地址}}, {{被告电话}}, {{被告公司名称}}, {{被告公司代码}}, {{被告法定代表人}}, {{被告法定代表人职务}}')
    print('案件相关: {{案由}}, {{诉讼请求1-5}}, {{事实与理由}}, {{证据清单}}, {{受诉法院}}, {{具状人}}, {{日期}}')
    print('代理人相关: {{代理人姓名}}, {{律所名称}}, {{代理人电话}}, {{代理人执业证号}}, {{委托权限}}')
