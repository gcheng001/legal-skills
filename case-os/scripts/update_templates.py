#!/usr/bin/env python3
"""
修改起诉状模板（移除证据清单），创建证据清单填空版模板
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def set_chinese_font(run, font_name='仿宋', font_size=14):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def create_complaint_template_without_evidence():
    """创建民事起诉状通用模板（不含证据清单）"""
    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # 设置中文字体
    style = doc.styles['Normal']
    style.font.name = '仿宋'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    style.font.size = Pt(14)

    # 标题
    title = doc.add_paragraph()
    title_run = title.add_run('民事起诉状')
    set_chinese_font(title_run, '黑体', 22)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(12)

    # 原告信息
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('原告：{{原告姓名}}')
    set_chinese_font(run, '黑体', 14)

    p = doc.add_paragraph()
    run = p.add_run('性别：{{原告性别}}　　民族：{{原告民族}}　　出生日期：{{原告出生日期}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('身份证号：{{原告身份证号}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('住所地：{{原告地址}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('联系电话：{{原告电话}}')
    set_chinese_font(run, '仿宋', 14)

    # 如果原告是公司
    p = doc.add_paragraph()
    run = p.add_run('（或）原告：{{原告公司名称}}')
    set_chinese_font(run, '黑体', 14)

    p = doc.add_paragraph()
    run = p.add_run('统一社会信用代码：{{原告公司代码}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('住所地：{{原告公司地址}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('法定代表人：{{原告法定代表人}}　　职务：{{原告法定代表人职务}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('联系电话：{{原告公司电话}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    # 被告信息
    p = doc.add_paragraph()
    run = p.add_run('被告：{{被告姓名}}')
    set_chinese_font(run, '黑体', 14)

    p = doc.add_paragraph()
    run = p.add_run('性别：{{被告性别}}　　民族：{{被告民族}}　　出生日期：{{被告出生日期}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('身份证号：{{被告身份证号}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('住所地：{{被告地址}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('联系电话：{{被告电话}}')
    set_chinese_font(run, '仿宋', 14)

    # 如果被告是公司
    p = doc.add_paragraph()
    run = p.add_run('（或）被告：{{被告公司名称}}')
    set_chinese_font(run, '黑体', 14)

    p = doc.add_paragraph()
    run = p.add_run('统一社会信用代码：{{被告公司代码}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('住所地：{{被告公司地址}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('法定代表人：{{被告法定代表人}}　　职务：{{被告法定代表人职务}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('联系电话：{{被告公司电话}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    # 案由
    p = doc.add_paragraph()
    run = p.add_run('案由：{{案由}}')
    set_chinese_font(run, '黑体', 14)

    doc.add_paragraph()

    # 诉讼请求
    p = doc.add_paragraph()
    run = p.add_run('诉讼请求：')
    set_chinese_font(run, '黑体', 14)

    p = doc.add_paragraph()
    run = p.add_run('1、{{诉讼请求1}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('2、{{诉讼请求2}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('3、{{诉讼请求3}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('4、{{诉讼请求4}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('5、{{诉讼请求5}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    # 事实与理由
    p = doc.add_paragraph()
    run = p.add_run('事实与理由：')
    set_chinese_font(run, '黑体', 14)

    p = doc.add_paragraph()
    run = p.add_run('{{事实与理由}}')
    set_chinese_font(run, '仿宋', 14)

    # 添加多行空白用于填写
    for _ in range(5):
        doc.add_paragraph()

    # 此致
    p = doc.add_paragraph()
    run = p.add_run('此致')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('{{受诉法院}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()
    doc.add_paragraph()

    # 具状人和日期（右对齐）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('具状人(签字、盖章)：')
    set_chinese_font(run, '仿宋', 14)
    run = p.add_run('{{具状人}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('日期：')
    set_chinese_font(run, '仿宋', 14)
    run = p.add_run('{{日期}}')
    set_chinese_font(run, '仿宋', 14)

    return doc

def create_evidence_list_template():
    """创建证据清单填空版模板"""
    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # 设置中文字体
    style = doc.styles['Normal']
    style.font.name = '仿宋'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    style.font.size = Pt(14)

    # 标题
    title = doc.add_paragraph()
    title_run = title.add_run('证 据 清 单')
    set_chinese_font(title_run, '黑体', 22)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(12)

    # 案件信息
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('案号：{{案号}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('案由：{{案由}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('原告：{{原告姓名}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('被告：{{被告姓名}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    # 提交人信息
    p = doc.add_paragraph()
    run = p.add_run('提交人：{{原告姓名}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('提交时间：{{提交时间}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    # 证据清单标题
    p = doc.add_paragraph()
    run = p.add_run('证据清单：')
    set_chinese_font(run, '黑体', 14)

    doc.add_paragraph()

    # 证据1
    p = doc.add_paragraph()
    run = p.add_run('证据1：')
    set_chinese_font(run, '黑体', 14)

    p = doc.add_paragraph()
    run = p.add_run('证据名称：{{证据1名称}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('证据来源：{{证据1来源}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('证明内容：{{证据1证明内容}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('页码/数量：{{证据1页码}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    # 证据2
    p = doc.add_paragraph()
    run = p.add_run('证据2：')
    set_chinese_font(run, '黑体', 14)

    p = doc.add_paragraph()
    run = p.add_run('证据名称：{{证据2名称}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('证据来源：{{证据2来源}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('证明内容：{{证据2证明内容}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('页码/数量：{{证据2页码}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    # 证据3
    p = doc.add_paragraph()
    run = p.add_run('证据3：')
    set_chinese_font(run, '黑体', 14)

    p = doc.add_paragraph()
    run = p.add_run('证据名称：{{证据3名称}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('证据来源：{{证据3来源}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('证明内容：{{证据3证明内容}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('页码/数量：{{证据3页码}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    # 证据4
    p = doc.add_paragraph()
    run = p.add_run('证据4：')
    set_chinese_font(run, '黑体', 14)

    p = doc.add_paragraph()
    run = p.add_run('证据名称：{{证据4名称}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('证据来源：{{证据4来源}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('证明内容：{{证据4证明内容}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('页码/数量：{{证据4页码}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    # 证据5
    p = doc.add_paragraph()
    run = p.add_run('证据5：')
    set_chinese_font(run, '黑体', 14)

    p = doc.add_paragraph()
    run = p.add_run('证据名称：{{证据5名称}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('证据来源：{{证据5来源}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('证明内容：{{证据5证明内容}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('页码/数量：{{证据5页码}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()
    doc.add_paragraph()

    # 签名区
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('提交人(签字/盖章)：')
    set_chinese_font(run, '仿宋', 14)
    run = p.add_run('{{提交人签字}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('日期：')
    set_chinese_font(run, '仿宋', 14)
    run = p.add_run('{{日期}}')
    set_chinese_font(run, '仿宋', 14)

    return doc

if __name__ == '__main__':
    base_path = '~/templates/civil'

    # 更新起诉状模板（不含证据清单）
    complaint_doc = create_complaint_template_without_evidence()
    complaint_path = os.path.join(base_path, '起诉状', '民事起诉状-通用模板.docx')
    complaint_doc.save(complaint_path)
    print(f'✓ 已更新起诉状模板（移除证据清单）: {complaint_path}')

    # 创建证据清单模板
    evidence_doc = create_evidence_list_template()
    evidence_path = os.path.join(base_path, '证据清单-填空模板.docx')
    evidence_doc.save(evidence_path)
    print(f'✓ 已创建证据清单模板: {evidence_path}')

    print('\n更新完成！')
    print('\n起诉状模板支持的占位符：')
    print('- 原告信息：{{原告姓名}}, {{原告性别}}, {{原告民族}}, {{原告出生日期}}, {{原告身份证号}}, {{原告地址}}, {{原告电话}}')
    print('- 原告公司信息：{{原告公司名称}}, {{原告公司代码}}, {{原告公司地址}}, {{原告法定代表人}}, {{原告法定代表人职务}}, {{原告公司电话}}')
    print('- 被告信息：{{被告姓名}}, {{被告性别}}, {{被告民族}}, {{被告出生日期}}, {{被告身份证号}}, {{被告地址}}, {{被告电话}}')
    print('- 被告公司信息：{{被告公司名称}}, {{被告公司代码}}, {{被告公司地址}}, {{被告法定代表人}}, {{被告法定代表人职务}}, {{被告公司电话}}')
    print('- 案件信息：{{案由}}, {{诉讼请求1-5}}, {{事实与理由}}, {{受诉法院}}, {{具状人}}, {{日期}}')

    print('\n证据清单模板支持的占位符：')
    print('- 案件信息：{{案号}}, {{案由}}, {{原告姓名}}, {{被告姓名}}')
    print('- 提交信息：{{提交时间}}, {{提交人签字}}, {{日期}}')
    print('- 证据1-5：{{证据X名称}}, {{证据X来源}}, {{证据X证明内容}}, {{证据X页码}}')
