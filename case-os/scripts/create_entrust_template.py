#!/usr/bin/env python3
"""
读取现有的委托手续模板，只添加需要填空的占位符
保留原有的律师信息，只填空：委托人、对方、案由、案号（可选）、时间、收费方式
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import re
import os

def set_chinese_font(run, font_name='仿宋', font_size=14):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def create_entrust_template_with_placeholders():
    """
    创建委托手续填空版模板
    基于标准格式，只填空关键信息
    """
    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # 设置中文字体
    style = doc.styles['Normal']
    style.font.name = '仿宋'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    style.font.size = Pt(14)

    # ========== 第一部分：委托合同 ==========
    title = doc.add_paragraph()
    title_run = title.add_run('委托代理合同')
    set_chinese_font(title_run, '黑体', 22)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(12)

    doc.add_paragraph()

    # 案号（可能为空）
    p = doc.add_paragraph()
    run = p.add_run('案号：{{案号}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    # 委托人（甲方）
    p = doc.add_paragraph()
    run = p.add_run('委托人（甲方）：{{委托人姓名}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('身份证号/统一社会信用代码：{{委托人身份证号}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('地址：{{委托人地址}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('电话：{{委托人电话}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    # 受托人（乙方）- 保留原有律师信息，不需要填空
    p = doc.add_paragraph()
    run = p.add_run('受托人（乙方）：浙江嘉瑞成律师事务所')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('地址：温州市鹿城区市府路525号同人恒玖大厦18楼')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('电话：0577-88999255')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    # 合同正文
    p = doc.add_paragraph()
    run = p.add_run('甲方因与{{对方姓名}}{{案由}}纠纷一案，委托乙方律师代理。经双方协商，订立本合同，共同遵照履行：')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    # 第一条
    p = doc.add_paragraph()
    run = p.add_run('一、乙方接受甲方委托，指派陈律师、李律师为甲方上述案件的一审/二审/执行阶段代理人。')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    # 第二条
    p = doc.add_paragraph()
    run = p.add_run('二、甲方委托乙方代理权限为：')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('（如：特别授权：代为起诉、应诉、承认、放弃、变更诉讼请求，进行和解、调解，代收法律文书等）')
    set_chinese_font(run, '仿宋', 12)

    doc.add_paragraph()

    # 第三条
    p = doc.add_paragraph()
    run = p.add_run('三、根据《浙江省律师服务收费标准》及双方协商，甲方向乙方缴纳律师费人民币{{收费金额}}元（{{收费方式}}）。')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('甲方应在本合同签订之日起{{付款期限}}日内付清上述费用。')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    # 第四条
    p = doc.add_paragraph()
    run = p.add_run('四、本合同有效期自签订之日起至本案本审终结止（判决、调解、撤诉或案外和解）。')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    # 第五条
    p = doc.add_paragraph()
    run = p.add_run('五、甲方应如实陈述案情，提供证据材料。如甲方捏造事实、伪造证据，乙方有权终止代理，所收费用不予退还。')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    # 第六条
    p = doc.add_paragraph()
    run = p.add_run('六、乙方应勤勉尽责，依法维护甲方合法权益。')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()
    doc.add_paragraph()

    # 签名区
    p = doc.add_paragraph()
    run = p.add_run('甲方（签字/盖章）：{{委托人签字}}')
    set_chinese_font(run, '仿宋', 14)
    run = p.add_run('　　　　　　　　')
    set_chinese_font(run, '仿宋', 14)
    run = p.add_run('乙方（盖章）：')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('日期：{{签订日期}}')
    set_chinese_font(run, '仿宋', 14)
    run = p.add_run('　　　　　　　　　　　　')
    set_chinese_font(run, '仿宋', 14)
    run = p.add_run('日期：{{签订日期}}')
    set_chinese_font(run, '仿宋', 14)

    # 分页
    doc.add_page_break()

    # ========== 第二部分：授权委托书 ==========
    title = doc.add_paragraph()
    title_run = title.add_run('授 权 委 托 书')
    set_chinese_font(title_run, '黑体', 22)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(12)

    doc.add_paragraph()

    # 案号（可能为空）
    p = doc.add_paragraph()
    run = p.add_run('案号：{{案号}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    # 委托人信息
    p = doc.add_paragraph()
    run = p.add_run('委托人：{{委托人姓名}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('身份证号/统一社会信用代码：{{委托人身份证号}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('住址/住所地：{{委托人地址}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('联系电话：{{委托人电话}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    # 受托人信息（固定律师信息）
    p = doc.add_paragraph()
    run = p.add_run('受委托人：陈律师')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('工作单位：浙江嘉瑞成律师事务所')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('职务：律师')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('联系电话：138XXXXXXXX')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('执业证号：133XXXXXXXXXXXXXX')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    # 委托事项
    p = doc.add_paragraph()
    run = p.add_run('现委托陈律师在{{委托人姓名}}与{{对方姓名}}{{案由}}纠纷一案中，作为我方委托代理人。')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('委托事项和权限如下：')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('{{委托权限}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('（如：特别授权：代为起诉、应诉、承认、放弃、变更诉讼请求，进行和解、调解，代收法律文书等）')
    set_chinese_font(run, '仿宋', 12)

    doc.add_paragraph()
    doc.add_paragraph()

    # 签名区
    p = doc.add_paragraph()
    run = p.add_run('委托人(签字/盖章)：{{委托人签字}}')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('日期：{{签订日期}}')
    set_chinese_font(run, '仿宋', 14)

    # 分页
    doc.add_page_break()

    # ========== 第三部分：律师事务所函 ==========
    title = doc.add_paragraph()
    title_run = title.add_run('律 师 事 务 所 函')
    set_chinese_font(title_run, '黑体', 22)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(12)

    doc.add_paragraph()

    # 案号（可能为空）
    p = doc.add_paragraph()
    run = p.add_run('案号：{{案号}}')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('{{受诉法院}}：')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('本所接受{{委托人姓名}}的委托，指派陈律师担任其与你院受理的（{{案由}}）纠纷一案中{{委托人姓名}}的委托代理人。')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('特此函告')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('附：1、授权委托书一份')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    run = p.add_run('　　2、律师证复印件一份')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()
    doc.add_paragraph()

    # 律所盖章区（右对齐）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('浙江嘉瑞成律师事务所（盖章）')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('日期：{{签订日期}}')
    set_chinese_font(run, '仿宋', 14)

    # 分页
    doc.add_page_break()

    # ========== 第四部分：介绍信 ==========
    title = doc.add_paragraph()
    title_run = title.add_run('介 绍 信')
    set_chinese_font(title_run, '黑体', 22)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('{{受诉法院}}：')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('兹介绍我所陈律师、李律师前往你处办理{{委托人姓名}}与{{对方姓名}}{{案由}}纠纷一案的相关事宜，请予接洽。')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('（有效期{{有效期限}}天）')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('浙江嘉瑞成律师事务所（盖章）')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('{{签订日期}}')
    set_chinese_font(run, '仿宋', 14)

    # 分页
    doc.add_page_break()

    # ========== 第五部分：所函（法院） ==========
    title = doc.add_paragraph()
    title_run = title.add_run('律 师 事 务 所 函')
    set_chinese_font(title_run, '黑体', 22)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('{{受诉法院}}：')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('兹指派我所陈律师、李律师在贵院受理的（{{案由}}）纠纷一案中担任{{委托人姓名}}的代理人。')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('特此函告')
    set_chinese_font(run, '仿宋', 14)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('浙江嘉瑞成律师事务所（盖章）')
    set_chinese_font(run, '仿宋', 14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('{{签订日期}}')
    set_chinese_font(run, '仿宋', 14)

    return doc

if __name__ == '__main__':
    base_path = '~/Codex/文件中转站/模板/民事案件'

    # 创建委托手续模板（使用固定律师信息）
    doc = create_entrust_template_with_placeholders()
    output_path = os.path.join(base_path, '委托手续-填空模板.docx')
    doc.save(output_path)
    print(f'✓ 已创建委托手续模板: {output_path}')

    print('\n委托手续模板说明：')
    print('- 律师信息已固定（浙江嘉瑞成律师事务所）')
    print('- 需要填空的字段：')
    print('  * {{案号}} - 可能为空')
    print('  * {{委托人姓名}}、{{委托人身份证号}}、{{委托人地址}}、{{委托人电话}}')
    print('  * {{对方姓名}}')
    print('  * {{案由}}')
    print('  * {{收费金额}}、{{收费方式}}、{{付款期限}}')
    print('  * {{委托权限}}')
    print('  * {{委托人签字}}、{{签订日期}}')
    print('  * {{受诉法院}}')
    print('  * {{有效期限}}')
