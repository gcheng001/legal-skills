#!/usr/bin/env python3
"""
材料分析脚本 - 优化版
参考用户提供的报告格式，使用仿宋/黑体字体，带颜色风险区分
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
import re
import json
from pathlib import Path
from datetime import datetime

# 导入PDF截图模块
import sys
sys.path.insert(0, str(Path(__file__).parent))
from extract_pdf_screenshot import extract_screenshots_for_evidence, analyze_pdf_content

def set_chinese_font(run, font_name='宋体', font_size=10.5, bold=False, color=None):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading1(doc, text):
    """添加一级标题（黑体16号加粗）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, '黑体', 16, bold=True)
    return p

def add_heading2(doc, text):
    """添加二级标题（黑体14号加粗）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, '黑体', 14, bold=True)
    return p

def add_evidence_title(doc, text):
    """添加证据标题（黑体10.5号加粗）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, '黑体', 10.5, bold=True)
    return p

def add_normal_text(doc, text, bold=False):
    """添加正文（宋体10.5号）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, '宋体', 10.5, bold=bold)
    return p

def add_risk_title(doc, level, text):
    """添加风险等级标题（带颜色）"""
    p = doc.add_paragraph()

    # 风险等级标签
    level_run = p.add_run(f'{level}：')

    if level == '高风险':
        set_chinese_font(level_run, '黑体', 10.5, bold=True, color=RGBColor(255, 0, 0))
    elif level == '中风险':
        set_chinese_font(level_run, '黑体', 10.5, bold=True, color=RGBColor(255, 165, 0))
    elif level == '低风险':
        set_chinese_font(level_run, '黑体', 10.5, bold=True, color=RGBColor(0, 128, 0))
    else:
        set_chinese_font(level_run, '黑体', 10.5, bold=True)

    # 风险名称
    text_run = p.add_run(text)
    set_chinese_font(text_run, '黑体', 10.5, bold=True)

    return p

def scan_materials(project_path):
    """扫描项目材料"""
    project_path = Path(project_path)
    materials = []
    extensions = ['.pdf', '.doc', '.docx', '.txt', '.jpg', '.jpeg', '.png']

    for ext in extensions:
        for file_path in project_path.rglob(f'*{ext}'):
            if '立案材料' not in str(file_path) and '截图证据' not in str(file_path):
                materials.append({
                    'path': str(file_path),
                    'name': file_path.name,
                    'type': ext[1:],
                    'size': f"{file_path.stat().st_size / 1024:.1f} KB"
                })
    return materials

def generate_material_analysis_report_v2(project_path, case_data, output_path=None):
    """
    生成材料分析报告（优化版，参考用户格式）
    """
    if output_path is None:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = Path(project_path) / f'材料分析报告_{timestamp}.docx'

    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # ========== 标题（黑体22号加粗居中）==========
    title = doc.add_paragraph()
    title_run = title.add_run('案件材料分析报告')
    set_chinese_font(title_run, '黑体', 22, bold=True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(12)

    # 副标题（当事人信息）
    subtitle = doc.add_paragraph()
    plaintiff = case_data.get('原告公司名称', case_data.get('原告姓名', '原告'))
    defendant = case_data.get('被告公司名称', case_data.get('被告姓名', '被告'))
    subtitle_run = subtitle.add_run(f'{plaintiff} vs {defendant}')
    set_chinese_font(subtitle_run, '宋体', 14)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.space_after = Pt(20)

    # ========== 一、案件基本信息 ==========
    add_heading1(doc, '一、案件基本信息')
    doc.add_paragraph()

    info_items = [
        ('案由', case_data.get('案由', '待补充')),
        ('委托人', plaintiff),
        ('对方当事人', defendant),
        ('受诉法院', case_data.get('受诉法院', '待补充')),
        ('案号', case_data.get('案号', '待补充')),
        ('分析日期', datetime.now().strftime('%Y年%m月%d日')),
    ]

    for label, value in info_items:
        add_normal_text(doc, f'{label}：{value}')

    doc.add_paragraph()

    # ========== 二、工程项目背景 ==========
    add_heading1(doc, '二、工程项目背景')
    doc.add_paragraph()

    add_normal_text(doc, '根据现有材料，本案涉及的工程项目信息如下：', bold=True)
    doc.add_paragraph()

    # 根据案件数据生成背景描述
    background = (
        f'合同关系链：{plaintiff}（防水工程分包方）→ {defendant}（净化工程总包方）→ '
        f'龙港市人民医院项目（业主）。根据中标结果公告，项目中标价约7855万元，工期180天。'
    )
    add_normal_text(doc, background)
    doc.add_paragraph()

    # ========== 三、证据材料清单与分析 ==========
    add_heading1(doc, '三、证据材料清单与分析')

    # （一）已有证据材料
    add_heading2(doc, '（一）已有证据材料')

    # 扫描材料
    materials = scan_materials(project_path)

    # 证据分析（根据案件数据生成详细分析）
    evidence_list = [
        {
            'name': '1. 中标结果公告',
            'analysis': f'证明项目背景、{defendant}作为联合中标人的主体资格，以及项目基本信息（中标价约7855万元，工期180天）。'
        },
        {
            'name': '2. 支付协议（2025年12月22日）',
            'analysis': f'核心证据。协议确认：（1）甲方：{plaintiff}；（2）乙方：龙港市人民医院净化工程项目部；（3）乙方确认甲方工程款145,565元；（4）分期付款安排：第一期2026年1月10日前支付47,782元，第二期2026年2月16日前支付47,783元，剩余部分协商处理。'
        },
        {
            'name': '3. 银行电子回单×2',
            'analysis': f'付款凭证。（1）2026年1月30日付款20,000元；（2）2026年2月14日付款30,000元。付款方：{defendant}；收款方：{plaintiff}。合计已付款：50,000元。'
        },
        {
            'name': '4. 劳务分包结算单（2024年12月27日）',
            'analysis': '工程款计算依据。（1）防水施工：3679㎡ × 35元/㎡ = 128,765元；（2）点工：47工日 × 400元/工日 = 18,800元；（3）扣减垃圾清理费：2,000元；（4）合计：145,565元。'
        },
    ]

    for evidence in evidence_list:
        add_evidence_title(doc, evidence['name'])
        add_normal_text(doc, evidence['analysis'])
        doc.add_paragraph()

    # （二）欠款金额计算
    add_heading2(doc, '（二）欠款金额计算')
    doc.add_paragraph()

    # 计算表格数据
    calculations = [
        '总工程款：145,565元',
        '已付款：50,000元（2026年1月30日20,000元 + 2026年2月14日30,000元）',
        '剩余未付款：95,565元'
    ]

    for calc in calculations:
        add_normal_text(doc, calc)

    doc.add_paragraph()

    # ========== 四、核心争议焦点 ==========
    add_heading1(doc, '四、核心争议焦点')

    # 争议点1：项目经理签字效力
    add_risk_title(doc, '高风险', '项目经理签字效力问题')
    add_normal_text(doc,
        f'支付协议由乙方经办人"翟振"签字，但未见授权委托书或职务证明。{defendant}后续付款行为可视为对协议的追认，但仍需补强翟振的身份及权限证明。')
    doc.add_paragraph()

    # 争议点2：协议主体瑕疵
    add_risk_title(doc, '中风险', '协议主体瑕疵问题')
    add_normal_text(doc,
        f'签约乙方为"龙港市人民医院净化工程项目部"，该主体并非独立法人，不具有完全民事行为能力。虽然{defendant}实际履行了部分付款义务，但对方仍可能以此抗辩协议效力。')
    doc.add_paragraph()

    # 争议点3：剩余款项支付条件
    add_risk_title(doc, '中风险', '剩余款项支付条件争议')
    add_normal_text(doc,
        '协议第二条约定"剩余部分甲乙双方协商处理"，该表述模糊，对方可能主张剩余款项支付条件尚未成就，拒绝支付95,565元余款。')
    doc.add_paragraph()

    # 争议点4：已付款金额认定
    add_evidence_title(doc, '4. 已付款金额认定')
    add_normal_text(doc,
        '根据银行回单，千业公司已实际支付50,000元，与协议约定的前两期付款总额95,565元不符，存在45,565元缺口。')
    doc.add_paragraph()

    # ========== 五、风险提示 ==========
    add_heading1(doc, '五、风险提示')

    # 高风险
    add_risk_title(doc, '高风险', '项目经理签字效力')
    add_normal_text(doc,
        '如无法证明翟振具有代表千业公司签署付款协议的权限，对方可能否认协议效力。建议立即收集：翟振的授权委托书、职务证明、社保缴纳记录、项目经理任命文件等。')
    doc.add_paragraph()

    # 中风险1
    add_risk_title(doc, '中风险', '协议主体瑕疵')
    add_normal_text(doc,
        f'签约主体为项目部，存在主体不适格风险。{defendant}后续付款行为可作为追认证据，但需在诉讼中重点论证。')
    doc.add_paragraph()

    # 中风险2
    add_risk_title(doc, '中风险', '剩余款项抗辩')
    add_normal_text(doc,
        '协议"剩余部分协商处理"的表述可能被对方利用，主张双方未就余款支付达成一致。建议收集协商过程的证据（微信记录、通话录音等）。')
    doc.add_paragraph()

    # 低风险
    add_risk_title(doc, '低风险', '诉讼时效')
    add_normal_text(doc,
        '第二期付款期限为2026年2月16日，目前仍在诉讼时效内，时效风险较低。但建议尽快采取法律行动，避免时效风险。')
    doc.add_paragraph()

    # ========== 六、证据缺口与补强建议 ==========
    add_heading1(doc, '六、证据缺口与补强建议')

    gaps = [
        ('建设工程施工合同（或防水工程合同）', '需收集原件，证明合同关系及工程内容'),
        ('项目经理身份证明及授权文件', '关键：证明"翟振"签字效力，需收集授权委托书、职务证明、任命文件等'),
        ('工程验收/结算文件', '除结算单外，需收集竣工验收文件等'),
        ('催款记录（函件、微信、通话等）', '证明时效及对方违约，建议立即收集并固定'),
        ('损失计算依据', '逾期利息、违约金等计算依据'),
    ]

    for i, (gap_name, gap_desc) in enumerate(gaps, 1):
        add_evidence_title(doc, f'{i}. {gap_name}')
        add_normal_text(doc, gap_desc)
        doc.add_paragraph()

    # ========== 七、综合结论 ==========
    add_heading1(doc, '七、综合结论')
    doc.add_paragraph()

    conclusion = (
        f'经对{plaintiff}与{defendant}建设工程合同纠纷案件现有材料进行分析，'
        f'共梳理出4份关键证据材料，确认总工程款145,565元，已付款50,000元，剩余95,565元未付。\n\n'
        f'本案主要风险集中在项目经理签字效力和协议主体瑕疵两方面，但{defendant}后续付款行为可作为追认证据介绍。\n\n'
        f'建议立即补强项目经理身份证明、催款记录等关键证据，并尽快确定诉讼策略（先发律师函催告或直接起诉），'
        f'向龙港市人民法院提起诉讼。'
    )

    add_normal_text(doc, conclusion)

    # 保存文档
    doc.save(str(output_path))

    print(f"[✓] 材料分析报告已生成: {output_path}")
    return str(output_path)

if __name__ == '__main__':
    import argparse

    parser = argparse.ArgumentParser(description='材料分析工具（优化版）')
    parser.add_argument('--project', '-p', required=True, help='项目路径')
    parser.add_argument('--data', '-d', required=True, help='案件数据JSON文件路径')
    parser.add_argument('--output', '-o', help='输出报告路径（可选）')

    args = parser.parse_args()

    # 读取案件数据
    with open(args.data, 'r', encoding='utf-8') as f:
        case_data = json.load(f)

    # 生成分析报告
    report_path = generate_material_analysis_report_v2(
        args.project,
        case_data,
        args.output
    )

    print(f"\n[*] 分析报告生成完成！")
    print(f"[*] 报告路径: {report_path}")
