#!/usr/bin/env python3
"""
构建通用答辩状模板
仅在目标文件缺失时创建；目标已存在时拒绝覆盖。
支持 --output 参数指定任意输出路径，以便在不碰 live 模板的情况下验证重建行为。
"""

import sys
import os
import argparse
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

TEMPLATE_DIR = "~/templates/civil/答辩状"
TARGET_FILE = os.path.join(TEMPLATE_DIR, "民事答辩状-通用模板.docx")


def set_chinese_font(run, font_name="仿宋", font_size=14):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def check_target_exists(path):
    """检查目标文件是否已存在。"""
    if os.path.exists(path):
        print(f"[拒绝覆盖] 目标文件已存在：{path}", file=sys.stderr)
        print("[拒绝覆盖] 如需重新生成，请先删除目标文件。", file=sys.stderr)
        return True
    return False


def ensure_dir(path):
    """确保目标目录存在。"""
    os.makedirs(os.path.dirname(path), exist_ok=True)


def create_defense_template():
    """创建民事答辩状通用模板。"""
    doc = Document()

    # 页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # 标题
    title = doc.add_paragraph()
    title_run = title.add_run("民事答辩状")
    set_chinese_font(title_run, "黑体", 22)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(20)

    # 答辩人（被告）信息
    p = doc.add_paragraph()
    run = p.add_run("答辩人（被告）：")
    set_chinese_font(run, "黑体", 14)
    run = p.add_run("{{原告姓名}}")
    set_chinese_font(run, "仿宋", 14)

    p = doc.add_paragraph()
    run = p.add_run("性别：")
    set_chinese_font(run, "仿宋", 14)
    run = p.add_run("{{原告性别}}")
    set_chinese_font(run, "仿宋", 14)
    run = p.add_run("　　民族：")
    set_chinese_font(run, "仿宋", 14)
    run = p.add_run("{{原告民族}}")
    set_chinese_font(run, "仿宋", 14)
    run = p.add_run("　　出生日期：")
    set_chinese_font(run, "仿宋", 14)
    run = p.add_run("{{原告出生日期}}")
    set_chinese_font(run, "仿宋", 14)

    p = doc.add_paragraph()
    run = p.add_run("身份证号：")
    set_chinese_font(run, "仿宋", 14)
    run = p.add_run("{{原告身份证号}}")
    set_chinese_font(run, "仿宋", 14)

    p = doc.add_paragraph()
    run = p.add_run("住所地：")
    set_chinese_font(run, "仿宋", 14)
    run = p.add_run("{{原告地址}}")
    set_chinese_font(run, "仿宋", 14)

    p = doc.add_paragraph()
    run = p.add_run("联系电话：")
    set_chinese_font(run, "仿宋", 14)
    run = p.add_run("{{原告电话}}")
    set_chinese_font(run, "仿宋", 14)

    p = doc.add_paragraph()
    run = p.add_run("（或）答辩人（被告）：")
    set_chinese_font(run, "黑体", 14)
    run = p.add_run("{{原告公司名称}}")
    set_chinese_font(run, "仿宋", 14)

    p = doc.add_paragraph()
    run = p.add_run("统一社会信用代码：")
    set_chinese_font(run, "仿宋", 14)
    run = p.add_run("{{原告公司代码}}")
    set_chinese_font(run, "仿宋", 14)

    p = doc.add_paragraph()
    run = p.add_run("住所地：")
    set_chinese_font(run, "仿宋", 14)
    run = p.add_run("{{原告公司地址}}")
    set_chinese_font(run, "仿宋", 14)

    p = doc.add_paragraph()
    run = p.add_run("法定代表人：")
    set_chinese_font(run, "仿宋", 14)
    run = p.add_run("{{原告法定代表人}}")
    set_chinese_font(run, "仿宋", 14)
    run = p.add_run("　　职务：")
    set_chinese_font(run, "仿宋", 14)
    run = p.add_run("{{原告法定代表人职务}}")
    set_chinese_font(run, "仿宋", 14)

    p = doc.add_paragraph()
    run = p.add_run("联系电话：")
    set_chinese_font(run, "仿宋", 14)
    run = p.add_run("{{原告公司电话}}")
    set_chinese_font(run, "仿宋", 14)

    doc.add_paragraph()

    # 被答辩人（原告）信息
    p = doc.add_paragraph()
    run = p.add_run("被答辩人（原告）：")
    set_chinese_font(run, "黑体", 14)
    run = p.add_run("{{被告姓名}}")
    set_chinese_font(run, "仿宋", 14)

    p = doc.add_paragraph()
    run = p.add_run("性别：")
    set_chinese_font(run, "仿宋", 14)
    run = p.add_run("{{被告性别}}")
    set_chinese_font(run, "仿宋", 14)
    run = p.add_run("　　民族：")
    set_chinese_font(run, "仿宋", 14)
    run = p.add_run("{{被告民族}}")
    set_chinese_font(run, "仿宋", 14)
    run = p.add_run("　　出生日期：")
    set_chinese_font(run, "仿宋", 14)
    run = p.add_run("{{被告出生日期}}")
    set_chinese_font(run, "仿宋", 14)

    p = doc.add_paragraph()
    run = p.add_run("身份证号：")
    set_chinese_font(run, "仿宋", 14)
    run = p.add_run("{{被告身份证号}}")
    set_chinese_font(run, "仿宋", 14)

    p = doc.add_paragraph()
    run = p.add_run("住所地：")
    set_chinese_font(run, "仿宋", 14)
    run = p.add_run("{{被告地址}}")
    set_chinese_font(run, "仿宋", 14)

    p = doc.add_paragraph()
    run = p.add_run("联系电话：")
    set_chinese_font(run, "仿宋", 14)
    run = p.add_run("{{被告电话}}")
    set_chinese_font(run, "仿宋", 14)

    p = doc.add_paragraph()
    run = p.add_run("（或）被答辩人（原告）：")
    set_chinese_font(run, "黑体", 14)
    run = p.add_run("{{被告公司名称}}")
    set_chinese_font(run, "仿宋", 14)

    p = doc.add_paragraph()
    run = p.add_run("统一社会信用代码：")
    set_chinese_font(run, "仿宋", 14)
    run = p.add_run("{{被告公司代码}}")
    set_chinese_font(run, "仿宋", 14)

    p = doc.add_paragraph()
    run = p.add_run("住所地：")
    set_chinese_font(run, "仿宋", 14)
    run = p.add_run("{{被告公司地址}}")
    set_chinese_font(run, "仿宋", 14)

    p = doc.add_paragraph()
    run = p.add_run("法定代表人：")
    set_chinese_font(run, "仿宋", 14)
    run = p.add_run("{{被告法定代表人}}")
    set_chinese_font(run, "仿宋", 14)
    run = p.add_run("　　职务：")
    set_chinese_font(run, "仿宋", 14)
    run = p.add_run("{{被告法定代表人职务}}")
    set_chinese_font(run, "仿宋", 14)

    p = doc.add_paragraph()
    run = p.add_run("联系电话：")
    set_chinese_font(run, "仿宋", 14)
    run = p.add_run("{{被告公司电话}}")
    set_chinese_font(run, "仿宋", 14)

    doc.add_paragraph()

    # 案由
    p = doc.add_paragraph()
    run = p.add_run("案由：")
    set_chinese_font(run, "黑体", 14)
    run = p.add_run("{{案由}}")
    set_chinese_font(run, "仿宋", 14)

    doc.add_paragraph()

    # 答辩请求
    p = doc.add_paragraph()
    run = p.add_run("答辩请求：")
    set_chinese_font(run, "黑体", 14)

    for i in range(1, 4):
        p = doc.add_paragraph()
        # 使用普通字符串拼接，避免 f-string 转义混淆；明确输出双花括号占位符
        run = p.add_run(str(i) + "、" + "{{答辩请求" + str(i) + "}}")
        set_chinese_font(run, "仿宋", 14)
        p.paragraph_format.first_line_indent = Cm(0.74)

    doc.add_paragraph()

    # 事实与理由（使用双花括号占位符，与 live 模板一致）
    p = doc.add_paragraph()
    run = p.add_run("事实与理由：")
    set_chinese_font(run, "黑体", 14)

    p = doc.add_paragraph()
    run = p.add_run("{{事实与理由}}")
    set_chinese_font(run, "仿宋", 14)

    doc.add_paragraph()

    # 受诉法院
    p = doc.add_paragraph()
    run = p.add_run("此致")
    set_chinese_font(run, "仿宋", 14)

    p = doc.add_paragraph()
    run = p.add_run("　　　　　　{{受诉法院}}")
    set_chinese_font(run, "仿宋", 14)

    doc.add_paragraph()
    doc.add_paragraph()

    # 签章与日期（与 live 模板完全一致）
    p = doc.add_paragraph()
    run = p.add_run("答辩人（签字、盖章）：{{具状人}}")
    set_chinese_font(run, "仿宋", 14)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("附：1. 答辩状副本　份")
    set_chinese_font(run, "仿宋", 14)

    p = doc.add_paragraph()
    run = p.add_run("　　　2. 证据材料　　份")
    set_chinese_font(run, "仿宋", 14)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("日期：{{日期}}")
    set_chinese_font(run, "仿宋", 14)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    return doc


def main():
    parser = argparse.ArgumentParser(description="构建通用答辩状模板")
    parser.add_argument(
        "--output", "-o",
        help=f"输出路径（默认：{TARGET_FILE}）",
        default=None
    )
    args = parser.parse_args()

    output_path = args.output if args.output else TARGET_FILE

    if check_target_exists(output_path):
        sys.exit(1)

    ensure_dir(output_path)
    doc = create_defense_template()
    doc.save(output_path)
    print(f"[创建完成] {output_path}")


if __name__ == "__main__":
    main()